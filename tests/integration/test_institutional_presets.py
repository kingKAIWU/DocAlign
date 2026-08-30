from __future__ import annotations

from pathlib import Path

import pytest
from docalign_core.docx.parser import parse_docx
from docalign_core.domain.document_ir import ParagraphIR
from docalign_core.domain.enums import SemanticRole
from docalign_core.domain.formatting_spec import (
    FormattingSpec,
    bigc_master_thesis_2025_reference_spec,
    gbt_9704_body_reference_spec,
    nankai_thesis_2026_reference_spec,
)
from docalign_core.services.processing import process_document
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

ROLE_BY_TEXT = {
    "规范验收标题": SemanticRole.TITLE,
    "摘要正文内容。": SemanticRole.ABSTRACT_BODY,
    "关键词：格式；验证": SemanticRole.KEYWORDS,
    "1 章标题": SemanticRole.HEADING_1,
    "1.1 一级标题": SemanticRole.HEADING_2,
    "1.1.1 二级标题": SemanticRole.HEADING_3,
    "1.1.1.1 三级标题": SemanticRole.HEADING_4,
    "正文中文与 English content。": SemanticRole.BODY,
    "图 1 处理流程": SemanticRole.FIGURE_CAPTION,
    "表 1 验收结果": SemanticRole.TABLE_CAPTION,
    "参考文献": SemanticRole.BIBLIOGRAPHY_HEADING,
    "[1] Example reference.": SemanticRole.BIBLIOGRAPHY_ENTRY,
    "附录 A 验收数据": SemanticRole.APPENDIX_HEADING,
    "附录正文内容。": SemanticRole.APPENDIX_BODY,
}


def _source_document(path: Path) -> None:
    document = Document()
    for text in ROLE_BY_TEXT:
        document.add_paragraph(text)
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "项目"
    table.cell(0, 1).text = "结果"
    table.cell(1, 0).text = "格式"
    table.cell(1, 1).text = "通过"
    document.save(path)


def _role_tagged_ir(path: Path):
    document_ir = parse_docx(path)
    for block in document_ir.blocks:
        if isinstance(block, ParagraphIR) and block.text in ROLE_BY_TEXT:
            block.detected_role = ROLE_BY_TEXT[block.text]
            block.role_confidence = 1
    return document_ir


def _paragraph(document: Document, text: str):
    return next(paragraph for paragraph in document.paragraphs if paragraph.text == text)


def _east_asia_font(paragraph) -> str | None:
    run_fonts = paragraph.runs[0]._r.get_or_add_rPr().find(qn("w:rFonts"))
    return run_fonts.get(qn("w:eastAsia")) if run_fonts is not None else None


def _assert_paragraph(
    document: Document,
    text: str,
    *,
    east_asia: str,
    size_pt: float,
    line_spacing_pt: float | None = None,
    alignment: WD_ALIGN_PARAGRAPH | None = None,
) -> None:
    paragraph = _paragraph(document, text)
    paragraph_format = paragraph.paragraph_format
    style_format = paragraph.style.paragraph_format
    assert _east_asia_font(paragraph) == east_asia
    assert paragraph.runs[0].font.size is not None
    assert paragraph.runs[0].font.size.pt == pytest.approx(size_pt, abs=0.05)
    if line_spacing_pt is not None:
        line_spacing = paragraph_format.line_spacing or style_format.line_spacing
        assert line_spacing is not None
        assert line_spacing.pt == pytest.approx(line_spacing_pt, abs=0.05)
    if alignment is not None:
        assert (paragraph.alignment or style_format.alignment) == alignment


@pytest.mark.parametrize(
    ("preset_id", "spec", "margins"),
    [
        (
            "gbt-9704-2012-body-reference-cn",
            gbt_9704_body_reference_spec(),
            (37, 35, 28, 26),
        ),
        (
            "nankai-thesis-2026-reference-cn",
            nankai_thesis_2026_reference_spec(),
            (38, 38, 32, 32),
        ),
        (
            "bigc-master-thesis-2025-reference-cn",
            bigc_master_thesis_2025_reference_spec(),
            (30, 25, 25, 25),
        ),
    ],
)
def test_institutional_reference_acceptance_fixture(
    preset_id: str,
    spec: FormattingSpec,
    margins: tuple[float, float, float, float],
    tmp_path: Path,
) -> None:
    source = tmp_path / f"{preset_id}-source.docx"
    output = tmp_path / f"{preset_id}-output.docx"
    _source_document(source)
    document_ir = _role_tagged_ir(source)

    result = process_document(
        source,
        document_ir,
        spec,
        output,
        job_id=f"acceptance-{preset_id}",
        artifact_dir=tmp_path / f"{preset_id}-artifacts",
    )

    assert result.audit.validation.valid
    assert parse_docx(output).content_fingerprint.digest == document_ir.content_fingerprint.digest
    formatted = Document(output)
    section = formatted.sections[0]
    actual_margins = (
        section.top_margin.mm,
        section.bottom_margin.mm,
        section.left_margin.mm,
        section.right_margin.mm,
    )
    assert actual_margins == pytest.approx(margins, abs=0.1)

    if preset_id == "gbt-9704-2012-body-reference-cn":
        _assert_paragraph(
            formatted,
            "规范验收标题",
            east_asia="方正小标宋简体",
            size_pt=22,
            line_spacing_pt=28.95,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
        )
        _assert_paragraph(
            formatted,
            "正文中文与 English content。",
            east_asia="仿宋_GB2312",
            size_pt=16,
            line_spacing_pt=28.95,
            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        )
        _assert_paragraph(formatted, "1 章标题", east_asia="黑体", size_pt=16)
        _assert_paragraph(formatted, "1.1 一级标题", east_asia="楷体_GB2312", size_pt=16)
        _assert_paragraph(formatted, "1.1.1 二级标题", east_asia="仿宋_GB2312", size_pt=16)
        body_paragraph = _paragraph(formatted, "正文中文与 English content。")
        first_line_indent = (
            body_paragraph.paragraph_format.first_line_indent
            or body_paragraph.style.paragraph_format.first_line_indent
        )
        assert first_line_indent is not None
        assert first_line_indent.pt == pytest.approx(32, abs=0.05)
        return

    _assert_paragraph(
        formatted,
        "正文中文与 English content。",
        east_asia="宋体",
        size_pt=12,
        line_spacing_pt=20,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )
    _assert_paragraph(formatted, "图 1 处理流程", east_asia="宋体", size_pt=10.5)
    _assert_paragraph(
        formatted,
        "[1] Example reference.",
        east_asia="宋体",
        size_pt=10.5,
        line_spacing_pt=16,
    )

    if preset_id == "nankai-thesis-2026-reference-cn":
        _assert_paragraph(
            formatted,
            "1 章标题",
            east_asia="黑体",
            size_pt=16,
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
        )
        _assert_paragraph(formatted, "1.1 一级标题", east_asia="黑体", size_pt=14)
        _assert_paragraph(formatted, "1.1.1 二级标题", east_asia="黑体", size_pt=13)
        _assert_paragraph(formatted, "1.1.1.1 三级标题", east_asia="黑体", size_pt=12)
    else:
        _assert_paragraph(
            formatted,
            "1 章标题",
            east_asia="黑体",
            size_pt=16,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
        )
        _assert_paragraph(formatted, "1.1 一级标题", east_asia="黑体", size_pt=15)
        _assert_paragraph(formatted, "1.1.1 二级标题", east_asia="黑体", size_pt=14)
