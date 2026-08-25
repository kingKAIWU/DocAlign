from __future__ import annotations

from pathlib import Path
from zipfile import ZipFile

from docalign_core.analysis.classifier import analyze_document
from docalign_core.docx.parser import parse_docx
from docalign_core.domain.document_ir import ParagraphIR, TableIR, UnsupportedBlockIR
from docalign_core.domain.enums import SemanticRole
from docalign_core.domain.formatting_spec import (
    DocumentFormattingSpec,
    FontSpec,
    FormattingSpec,
    LineSpacingMode,
    LineSpacingSpec,
    PageFormattingSpec,
    PageSize,
    ParagraphSpec,
    RoleFormattingSpec,
    VisualCleanupSpec,
    default_academic_spec,
    default_cleanup_spec,
)
from docalign_core.services.processing import process_document
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor
from lxml import etree


def test_end_to_end_format_preserves_content_and_is_idempotent(
    academic_docx: Path, tmp_path: Path
) -> None:
    spec = default_academic_spec()
    analysis = analyze_document(parse_docx(academic_docx))
    first_output = tmp_path / "formatted.docx"
    first = process_document(
        academic_docx,
        analysis.document_ir,
        spec,
        first_output,
        job_id="job_first",
        artifact_dir=tmp_path / "first-artifacts",
    )
    assert first.audit.validation.valid
    assert first_output.exists()
    first_ir = parse_docx(first_output)
    assert (
        first_ir.content_fingerprint.paragraph_texts
        == analysis.document_ir.content_fingerprint.paragraph_texts
    )
    assert (
        first_ir.content_fingerprint.table_cell_texts
        == analysis.document_ir.content_fingerprint.table_cell_texts
    )
    assert (
        first_ir.content_fingerprint.image_hashes
        == analysis.document_ir.content_fingerprint.image_hashes
    )

    second_analysis = analyze_document(first_ir)
    second_output = tmp_path / "formatted-twice.docx"
    second = process_document(
        first_output,
        second_analysis.document_ir,
        spec,
        second_output,
        job_id="job_second",
        artifact_dir=tmp_path / "second-artifacts",
    )
    assert second.audit.validation.valid
    assert (
        parse_docx(second_output).content_fingerprint.digest == first_ir.content_fingerprint.digest
    )
    assert second.audit.summary.changed_mutations == 0


def test_comprehensive_golden_preserves_protected_structures(
    comprehensive_docx: Path, tmp_path: Path
) -> None:
    source_ir = parse_docx(comprehensive_docx)
    paragraphs = [block for block in source_ir.blocks if isinstance(block, ParagraphIR)]
    tables = [block for block in source_ir.blocks if isinstance(block, TableIR)]
    unsupported = [block for block in source_ir.blocks if isinstance(block, UnsupportedBlockIR)]
    assert any(block.contains_hyperlink for block in paragraphs)
    assert any(block.contains_bookmark for block in paragraphs)
    assert any(block.contains_equation for block in paragraphs)
    assert any(block.contains_field for block in paragraphs)
    assert any(block.nested_tables_present for block in tables)
    assert unsupported and unsupported[0].text_preview == "受保护的未知顶层 OOXML 内容"
    assert "research_anchor" in source_ir.content_fingerprint.bookmark_names
    assert any("hyperlink" in item.relationship_type for item in source_ir.relationships)

    analysis = analyze_document(source_ir)
    output = tmp_path / "comprehensive-formatted.docx"
    result = process_document(
        comprehensive_docx,
        analysis.document_ir,
        default_academic_spec(),
        output,
        job_id="job-comprehensive",
        artifact_dir=tmp_path / "comprehensive-artifacts",
    )
    assert result.audit.validation.valid
    output_ir = parse_docx(output)
    assert output_ir.content_fingerprint.bookmark_names == (
        source_ir.content_fingerprint.bookmark_names
    )
    assert output_ir.content_fingerprint.relationship_signatures == (
        source_ir.content_fingerprint.relationship_signatures
    )
    assert output_ir.content_fingerprint.unsupported_block_signatures == (
        source_ir.content_fingerprint.unsupported_block_signatures
    )


def test_scoped_body_spec_preserves_page_headings_and_list_geometry(tmp_path: Path) -> None:
    source = tmp_path / "scoped-source.docx"
    document = Document()
    section = document.sections[0]
    section.left_margin = Mm(33)
    section.right_margin = Mm(27)
    heading = document.add_paragraph("1 研究背景")
    heading.style = document.styles["Normal"]
    heading_run = heading.runs[0]
    heading_run.bold = True
    heading_run.font.size = None
    document.add_paragraph("正文 Chinese and English mixed content。")
    list_paragraph = document.add_paragraph("保留列表缩进", style="List Bullet")
    source_list_style = list_paragraph.style.name
    document.save(source)

    analysis = analyze_document(parse_docx(source))
    spec = FormattingSpec(
        roles={
            SemanticRole.BODY: RoleFormattingSpec(
                font=FontSpec(
                    east_asia="宋体",
                    ascii="Times New Roman",
                    high_ansi="Times New Roman",
                    complex_script="Times New Roman",
                    size_pt=12,
                ),
                paragraph=ParagraphSpec(
                    line_spacing=LineSpacingSpec(mode=LineSpacingMode.MULTIPLE, value=1.5),
                    first_line_indent_pt=24,
                ),
            )
        }
    )
    output = tmp_path / "scoped-output.docx"
    result = process_document(
        source,
        analysis.document_ir,
        spec,
        output,
        job_id="job-scoped",
        artifact_dir=tmp_path / "scoped-artifacts",
    )
    assert result.audit.validation.valid

    formatted = Document(output)
    assert formatted.sections[0].left_margin == section.left_margin
    assert formatted.sections[0].right_margin == section.right_margin
    assert formatted.paragraphs[0].style.name == "Normal"
    assert formatted.paragraphs[1].style.name == "DA Body"
    assert formatted.paragraphs[2].style.name == source_list_style
    assert formatted.paragraphs[2].paragraph_format.first_line_indent is None
    list_run_fonts = formatted.paragraphs[2].runs[0]._r.rPr.find(qn("w:rFonts"))
    assert list_run_fonts is not None
    assert list_run_fonts.get(qn("w:eastAsia")) == "宋体"
    assert any(warning.code == "NUMBERING_LAYOUT_PRESERVED" for warning in result.audit.warnings)


def test_partial_page_spec_changes_size_without_inventing_margins(tmp_path: Path) -> None:
    source = tmp_path / "partial-page-source.docx"
    document = Document()
    section = document.sections[0]
    section.left_margin = Mm(34)
    section.right_margin = Mm(21)
    section.top_margin = Mm(19)
    section.bottom_margin = Mm(23)
    document.add_paragraph("正文内容")
    document.save(source)

    analysis = analyze_document(parse_docx(source))
    spec = FormattingSpec(
        document=DocumentFormattingSpec(page=PageFormattingSpec(size=PageSize.A4))
    )
    output = tmp_path / "partial-page-output.docx"
    result = process_document(
        source,
        analysis.document_ir,
        spec,
        output,
        job_id="job-partial-page",
        artifact_dir=tmp_path / "partial-page-artifacts",
    )

    assert result.audit.validation.valid
    formatted = Document(output).sections[0]
    assert formatted.left_margin == section.left_margin
    assert formatted.right_margin == section.right_margin
    assert formatted.top_margin == section.top_margin
    assert formatted.bottom_margin == section.bottom_margin


def test_baseline_formats_all_roles_and_heading_override_wins(tmp_path: Path) -> None:
    source = tmp_path / "baseline-source.docx"
    document = Document()
    document.add_paragraph("1 标题")
    document.add_paragraph("正文 English content")
    document.save(source)

    analysis = analyze_document(parse_docx(source))
    spec = FormattingSpec(
        baseline=RoleFormattingSpec(
            font=FontSpec(east_asia="宋体", ascii="Times New Roman", size_pt=12),
            paragraph=ParagraphSpec(
                line_spacing=LineSpacingSpec(mode=LineSpacingMode.MULTIPLE, value=1.5)
            ),
        ),
        roles={
            SemanticRole.HEADING_1: RoleFormattingSpec(
                font=FontSpec(east_asia="黑体", size_pt=16)
            )
        },
    )
    output = tmp_path / "baseline-output.docx"
    result = process_document(
        source,
        analysis.document_ir,
        spec,
        output,
        job_id="job-baseline",
        artifact_dir=tmp_path / "baseline-artifacts",
    )

    assert result.audit.validation.valid
    formatted = Document(output)
    heading_fonts = formatted.paragraphs[0].runs[0]._r.rPr.find(qn("w:rFonts"))
    body_fonts = formatted.paragraphs[1].runs[0]._r.rPr.find(qn("w:rFonts"))
    assert heading_fonts is not None
    assert heading_fonts.get(qn("w:eastAsia")) == "黑体"
    assert heading_fonts.get(qn("w:ascii")) == "Times New Roman"
    assert body_fonts is not None
    assert body_fonts.get(qn("w:eastAsia")) == "宋体"
    assert body_fonts.get(qn("w:ascii")) == "Times New Roman"


def test_visual_cleanup_is_complete_content_safe_and_idempotent(tmp_path: Path) -> None:
    source = tmp_path / "visual-cleanup-source.docx"
    document = Document()
    paragraph = document.add_paragraph()
    run = paragraph.add_run("红色高亮正文")
    run.font.color.rgb = RGBColor(0xEE, 0x00, 0x00)
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    run_shading = OxmlElement("w:shd")
    run_shading.set(qn("w:fill"), "00FF00")
    run._r.get_or_add_rPr().append(run_shading)
    paragraph_shading = OxmlElement("w:shd")
    paragraph_shading.set(qn("w:fill"), "CCCCCC")
    paragraph._p.get_or_add_pPr().append(paragraph_shading)

    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "单元格底纹文字"
    cell_shading = OxmlElement("w:shd")
    cell_shading.set(qn("w:fill"), "FFFF00")
    table.cell(0, 0)._tc.get_or_add_tcPr().append(cell_shading)

    background = OxmlElement("w:background")
    background.set(qn("w:color"), "E6E6E6")
    document._element.insert(0, background)
    document.save(source)

    source_ir = parse_docx(source)
    analysis = analyze_document(source_ir)
    spec = FormattingSpec(
        visual_cleanup=VisualCleanupSpec(
            text_color_hex="000000",
            remove_text_highlight=True,
            remove_character_shading=True,
            remove_paragraph_shading=True,
            remove_table_cell_shading=True,
            remove_page_background=True,
        )
    )
    first_output = tmp_path / "visual-cleanup-output.docx"
    first = process_document(
        source,
        analysis.document_ir,
        spec,
        first_output,
        job_id="job-visual-cleanup-first",
        artifact_dir=tmp_path / "visual-cleanup-first-artifacts",
    )

    assert first.audit.validation.valid
    first_ir = parse_docx(first_output)
    assert first_ir.content_fingerprint.digest == source_ir.content_fingerprint.digest
    _assert_visual_cleanup_xml(first_output)

    second_output = tmp_path / "visual-cleanup-output-twice.docx"
    second_ir = parse_docx(first_output)
    second = process_document(
        first_output,
        analyze_document(second_ir).document_ir,
        spec,
        second_output,
        job_id="job-visual-cleanup-second",
        artifact_dir=tmp_path / "visual-cleanup-second-artifacts",
    )
    assert second.audit.validation.valid
    assert second.audit.summary.changed_mutations == 0
    _assert_visual_cleanup_xml(second_output)


def test_default_cleanup_normalizes_emphasis_sections_and_wide_tables(
    tmp_path: Path,
) -> None:
    source = tmp_path / "default-cleanup-source.docx"
    document = Document()
    document.sections[0].header.paragraphs[0].add_run("彩色页眉").bold = True
    document.add_paragraph("1 常规整理")
    body = document.add_paragraph()
    body_run = body.add_run("这是一段需要清理直接格式的常规正文。")
    body_run.bold = True
    body_run.italic = True
    body_run.underline = True
    body_run.font.color.rgb = RGBColor(0xEE, 0x00, 0x00)
    body_run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    landscape = document.add_section(WD_SECTION.NEW_PAGE)
    landscape.orientation = WD_ORIENT.LANDSCAPE
    landscape.page_width, landscape.page_height = (
        landscape.page_height,
        landscape.page_width,
    )
    table = document.add_table(rows=2, cols=8)
    for row_index, row in enumerate(table.rows):
        for column_index, cell in enumerate(row.cells):
            cell_run = cell.paragraphs[0].add_run(f"R{row_index + 1}C{column_index + 1}")
            cell_run.bold = True
            cell_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    document.save(source)

    analysis = analyze_document(parse_docx(source))
    output = tmp_path / "default-cleanup-output.docx"
    result = process_document(
        source,
        analysis.document_ir,
        default_cleanup_spec(),
        output,
        job_id="job-default-cleanup",
        artifact_dir=tmp_path / "default-cleanup-artifacts",
    )

    assert result.audit.validation.valid
    formatted = Document(output)
    assert all(
        section.orientation == WD_ORIENT.PORTRAIT
        and section.page_width < section.page_height
        for section in formatted.sections
    )
    normalized_body = next(
        paragraph
        for paragraph in formatted.paragraphs
        if paragraph.text == "这是一段需要清理直接格式的常规正文。"
    )
    normalized_run = normalized_body.runs[0]
    assert normalized_run.bold is False
    assert normalized_run.italic is False
    assert normalized_run.underline is False
    fonts = normalized_run._r.rPr.find(qn("w:rFonts"))
    assert fonts is not None
    assert fonts.get(qn("w:eastAsia")) == "宋体"

    pagination_tags = ("keepNext", "keepLines", "pageBreakBefore")

    def pagination_flag_is_active(element: object) -> bool:
        value = element.get(qn("w:val"))
        return value not in {"0", "false", "off"}

    for paragraph in formatted.paragraphs:
        paragraph_properties = paragraph._p.pPr
        if paragraph_properties is not None:
            assert not any(
                pagination_flag_is_active(element)
                for tag in pagination_tags
                for element in paragraph_properties.findall(qn(f"w:{tag}"))
            )
        if paragraph.style.name.startswith("DA "):
            style_properties = paragraph.style.element.pPr
            if style_properties is not None:
                assert not any(
                    pagination_flag_is_active(element)
                    for tag in pagination_tags
                    for element in style_properties.findall(qn(f"w:{tag}"))
                )

    expected_table_size = Pt(9)
    table_runs = [
        run
        for row in formatted.tables[0].rows
        for cell in row.cells
        for paragraph in cell.paragraphs
        for run in paragraph.runs
        if run.text
    ]
    assert len(table_runs) == 16
    for run in table_runs:
        run_fonts = run._r.rPr.find(qn("w:rFonts"))
        assert run_fonts is not None
        assert run_fonts.get(qn("w:eastAsia")) == "宋体"
        assert run_fonts.get(qn("w:ascii")) == "Times New Roman"
        assert run.font.size == expected_table_size
        assert run.bold is False

    printable_width = int(
        formatted.sections[0].page_width.twips
        - formatted.sections[0].left_margin.twips
        - formatted.sections[0].right_margin.twips
    )
    grid_widths = [
        int(column.get(qn("w:w")))
        for column in formatted.tables[0]._tbl.tblGrid.findall(qn("w:gridCol"))
    ]
    assert sum(grid_widths) == printable_width
    assert formatted.tables[0].cell(0, 0).paragraphs[0].runs[0].bold is False
    assert formatted.sections[0].header.paragraphs[0].runs[0].bold is False
    _assert_visual_cleanup_xml(output)


def _assert_visual_cleanup_xml(path: Path) -> None:
    with ZipFile(path) as package:
        for name in package.namelist():
            if not name.startswith("word/") or not name.endswith(".xml"):
                continue
            root = etree.fromstring(package.read(name))
            assert not list(root.iter(qn("w:highlight")))
            assert not list(root.iter(qn("w:background")))
            for parent_name in ("w:rPr", "w:pPr", "w:tcPr"):
                assert all(
                    parent.find(qn("w:shd")) is None for parent in root.iter(qn(parent_name))
                )
            for run in root.iter(qn("w:r")):
                if not any((text.text or "") for text in run.iter(qn("w:t"))):
                    continue
                properties = run.find(qn("w:rPr"))
                assert properties is not None
                color = properties.find(qn("w:color"))
                assert color is not None
                assert color.get(qn("w:val")) == "000000"
                assert color.get(qn("w:themeColor")) is None
