from __future__ import annotations

import argparse
import json
from pathlib import Path
from zipfile import ZipFile

from docalign_core.analysis.classifier import analyze_document
from docalign_core.docx.parser import parse_docx
from docalign_core.domain.audit import CONTENT_INTEGRITY_CODES
from docalign_core.domain.document_ir import ParagraphIR
from docalign_core.domain.formatting_spec import cleanup_preset_catalog
from docalign_core.services.processing import process_document
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from lxml import etree

OUTPUT_LABELS = {
    "01-government-notice": "01-政府通知",
    "02-academic-paper": "02-学术论文",
    "03-business-report": "03-经营分析报告",
    "04-legal-contract": "04-软件服务合同",
    "05-meeting-minutes": "05-会议纪要",
    "06-resume": "06-个人简历",
    "07-training-manual": "07-操作手册",
    "08-financial-statement": "08-财务宽表",
}

def _visual_pollution(path: Path) -> dict[str, int]:
    totals = {
        "highlight": 0,
        "character_shading": 0,
        "paragraph_shading": 0,
        "cell_shading": 0,
        "page_background": 0,
    }
    with ZipFile(path) as package:
        for name in package.namelist():
            if not name.startswith("word/") or not name.endswith(".xml"):
                continue
            root = etree.fromstring(package.read(name))
            totals["highlight"] += sum(1 for _ in root.iter(qn("w:highlight")))
            totals["page_background"] += sum(1 for _ in root.iter(qn("w:background")))
            for run_properties in root.iter(qn("w:rPr")):
                totals["character_shading"] += int(run_properties.find(qn("w:shd")) is not None)
            for paragraph_properties in root.iter(qn("w:pPr")):
                totals["paragraph_shading"] += int(
                    paragraph_properties.find(qn("w:shd")) is not None
                )
            for cell_properties in root.iter(qn("w:tcPr")):
                totals["cell_shading"] += int(cell_properties.find(qn("w:shd")) is not None)
    return totals


def run(
    source_dir: Path,
    output_dir: Path,
    *,
    east_asian_font: str | None = None,
    all_font: str | None = None,
) -> list[dict[str, object]]:
    output_dir.mkdir(parents=True, exist_ok=True)
    reports: list[dict[str, object]] = []
    catalog = cleanup_preset_catalog()
    for source in sorted(source_dir.glob("*.docx")):
        source_ir = parse_docx(source)
        analysis = analyze_document(source_ir)
        document_kind = (
            analysis.summary.document_kind.value if analysis.summary.document_kind else "other"
        )
        has_wide_table = any(
            getattr(block, "columns_estimate", 0) >= 8 for block in analysis.document_ir.blocks
        )
        selected = (
            next(item for item in catalog if item.preset_id == "wide-table-clean-cn")
            if has_wide_table
            else next(
                (item for item in catalog if document_kind in item.recommended_kinds),
                catalog[0],
            )
        )
        output_label = OUTPUT_LABELS.get(source.stem, source.stem)
        output = output_dir / f"{output_label}-{selected.name}.docx"
        spec = selected.spec.model_copy(deep=True)
        if east_asian_font:
            if spec.baseline and spec.baseline.font:
                spec.baseline.font.east_asia = east_asian_font
            for role in spec.roles.values():
                if role.font:
                    role.font.east_asia = east_asian_font
            for target in (spec.tables, spec.headers, spec.footers):
                if target and target.font:
                    target.font.east_asia = east_asian_font
        if all_font:
            configured_fonts = []
            if spec.baseline and spec.baseline.font:
                configured_fonts.append(spec.baseline.font)
            configured_fonts.extend(role.font for role in spec.roles.values() if role.font)
            configured_fonts.extend(
                target.font
                for target in (spec.tables, spec.headers, spec.footers)
                if target and target.font
            )
            for font in configured_fonts:
                font.east_asia = all_font
                font.ascii = all_font
                font.high_ansi = all_font
                font.complex_script = all_font
        result = process_document(
            source,
            analysis.document_ir,
            spec,
            output,
            job_id=f"corpus-{source.stem}",
            artifact_dir=output_dir / f"audit-{source.stem}",
        )
        output_ir = parse_docx(output)
        rendered = Document(output)
        paragraphs = [
            block for block in analysis.document_ir.blocks if isinstance(block, ParagraphIR)
        ]
        reports.append(
            {
                "source": source.name,
                "output": output.name,
                "source_pollution": _visual_pollution(source),
                "output_pollution": _visual_pollution(output),
                "content_preserved": not any(
                    issue.code in CONTENT_INTEGRITY_CODES
                    for issue in result.audit.validation.issues
                ),
                "exact_fingerprint_preserved": (
                    source_ir.content_fingerprint.digest == output_ir.content_fingerprint.digest
                ),
                "validation_valid": result.audit.validation.valid,
                "validation_issues": len(result.audit.validation.issues),
                "all_sections_portrait": all(
                    section.orientation == WD_ORIENT.PORTRAIT
                    and section.page_width < section.page_height
                    for section in rendered.sections
                ),
                "document_kind": document_kind,
                "preset_id": selected.preset_id,
                "sections": len(rendered.sections),
                "paragraphs": len(paragraphs),
                "tables": analysis.summary.table_count,
                "images": analysis.summary.image_count,
                "role_counts": analysis.summary.role_counts,
                "reviewable_unknowns": analysis.summary.unknown_count,
                "warnings": [warning.code for warning in analysis.warnings],
                "auto_layout_splits": result.audit.summary.auto_layout_splits,
                "changed_mutations": result.audit.summary.changed_mutations,
                "validation_failures": result.audit.summary.validation_failures,
            }
        )
    (output_dir / "corpus-report.json").write_text(
        json.dumps(reports, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    return reports


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--source-dir", type=Path, required=True)
    parser.add_argument("--output-dir", type=Path, required=True)
    parser.add_argument("--east-asian-font")
    parser.add_argument("--all-font")
    args = parser.parse_args()
    for report in run(
        args.source_dir,
        args.output_dir,
        east_asian_font=args.east_asian_font,
        all_font=args.all_font,
    ):
        print(json.dumps(report, ensure_ascii=False))


if __name__ == "__main__":
    main()
