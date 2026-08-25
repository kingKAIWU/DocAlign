from __future__ import annotations

import argparse
import base64
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor

PNG_1X1 = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
)


def _shade(element, fill: str) -> None:
    properties = element.get_or_add_rPr() if element.tag == qn("w:r") else None
    if element.tag == qn("w:p"):
        properties = element.get_or_add_pPr()
    elif element.tag == qn("w:tc"):
        properties = element.get_or_add_tcPr()
    if properties is None:
        return
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def _pollute_run(run, *, color: str = "C00000", highlight=WD_COLOR_INDEX.YELLOW) -> None:
    run.bold = True
    run.italic = True
    run.underline = True
    run.font.name = "Arial"
    run.font.size = Pt(15)
    run.font.color.rgb = RGBColor.from_string(color)
    run.font.highlight_color = highlight
    _shade(run._r, "D9EAD3")


def _add_page_background(document: Document, color: str = "F2F2F2") -> None:
    background = OxmlElement("w:background")
    background.set(qn("w:color"), color)
    document._element.insert(0, background)


def _add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instruction, separate, text, end):
        run._r.append(node)


def _add_common_noise(document: Document, label: str) -> None:
    _add_page_background(document)
    header = document.sections[0].header.paragraphs[0]
    noisy_header = header.add_run(f"{label} · 内部资料")
    _pollute_run(noisy_header, color="7030A0", highlight=WD_COLOR_INDEX.TURQUOISE)
    _add_page_number(document.sections[0].footer.paragraphs[0])


def _title(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    _pollute_run(run)
    _shade(paragraph._p, "FFF2CC")


def _body(document: Document, text: str, *, noisy: bool = True) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    if noisy:
        _pollute_run(run, color="0070C0", highlight=WD_COLOR_INDEX.BRIGHT_GREEN)


def _landscape_section(document: Document):
    section = document.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Mm(12)
    section.right_margin = Mm(12)
    return section


def _fill_table_noise(table) -> None:
    for row_index, row in enumerate(table.rows):
        for column_index, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _shade(cell._tc, "FFF2CC" if row_index == 0 else "DDEBF7")
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    run.font.size = Pt(13 if column_index % 2 == 0 else 9)


def government_notice(path: Path) -> None:
    document = Document()
    _add_common_noise(document, "行政公文")
    _title(document, "关于开展 2026 年度信息安全检查的通知")
    number = document.add_paragraph("示例办发〔2026〕18号")
    number.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph("各部门、各直属单位：")
    _body(
        document,
        "为进一步落实信息安全责任，现组织开展年度检查。本段包含 intentionally mixed English，"
        "用于验证中西文字体、首行缩进和段落节奏。",
    )
    document.add_paragraph("一、检查范围")
    _body(document, "检查办公终端、业务系统、数据备份与外部协作平台的账号权限。")
    document.add_paragraph("（一）自查阶段")
    _body(document, "各单位应于九月十日前完成自查并提交问题清单。")
    document.add_paragraph("二、工作要求")
    for text in ("压实主体责任", "如实记录问题", "按期完成整改"):
        document.add_paragraph(text, style="List Bullet")
    table = document.add_table(rows=4, cols=4)
    for index, value in enumerate(("阶段", "时间", "责任单位", "交付物")):
        table.cell(0, index).text = value
    for row, values in enumerate(
        (
            ("自查", "9月1日-10日", "各部门", "自查表"),
            ("抽查", "9月11日-20日", "检查组", "记录单"),
            ("整改", "9月30日前", "责任单位", "整改报告"),
        ),
        start=1,
    ):
        for column, value in enumerate(values):
            table.cell(row, column).text = value
    _fill_table_noise(table)
    signature = document.add_paragraph("示例市数字化办公室\n2026年8月25日")
    signature.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    document.save(path)


def academic_paper(path: Path, image_path: Path) -> None:
    document = Document()
    _add_common_noise(document, "学术论文")
    _title(document, "面向复杂 Word 文档的确定性排版方法研究")
    author = document.add_paragraph("张三1，李四2")
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph("摘要")
    _body(document, "本文研究跨领域文档的结构识别、格式归一与内容完整性验证方法。")
    document.add_paragraph("关键词：Word；OOXML；自动排版；内容保护")
    document.add_paragraph("1 引言")
    long_paragraph = document.add_paragraph()
    first = long_paragraph.add_run(
        "实际文档常包含样式漂移、直接格式、突出显示以及多节页面设置。"
        "传统的全选后统一字体无法可靠恢复语义层级，也容易破坏字段、超链接和公式。"
    )
    _pollute_run(first)
    first.add_break(WD_BREAK.LINE)
    long_paragraph.add_run(
        "因此需要先识别标题、正文、题注与参考文献，再执行可审计的确定性变换。"
        "这一段故意使用手动换行，验证自动分段不会损坏内容。"
    )
    document.add_paragraph("1.1 研究方法")
    _body(document, "方法由结构解析、角色推断、格式计划、落盘验证和自动修复五个阶段组成。")
    picture = document.add_paragraph()
    picture.alignment = WD_ALIGN_PARAGRAPH.LEFT
    picture.add_run().add_picture(str(image_path), width=Cm(4.5), height=Cm(2.2))
    document.add_paragraph("图 1 确定性处理流程示意")
    document.add_paragraph("表 1 评价指标")
    table = document.add_table(rows=4, cols=5)
    headers = ("指标", "定义", "目标", "当前", "判定")
    for column, value in enumerate(headers):
        table.cell(0, column).text = value
    rows = (
        ("内容完整性", "文本、表格、图片指纹一致", "100%", "100%", "通过"),
        ("格式合规", "强制属性全部命中", "100%", "100%", "通过"),
        ("处理时延", "本地端到端耗时", "<30秒", "3秒", "通过"),
    )
    for row, values in enumerate(rows, start=1):
        for column, value in enumerate(values):
            table.cell(row, column).text = value
    _fill_table_noise(table)
    document.add_paragraph("2 结果与讨论")
    _body(document, "结果表明，内容安全约束与格式归一可以同时实现。")
    document.add_paragraph("参考文献")
    document.add_paragraph("[1] ECMA International. Office Open XML File Formats.")
    document.save(path)


def business_report(path: Path) -> None:
    document = Document()
    _add_common_noise(document, "经营分析")
    _title(document, "华东区域 2026 年第二季度经营分析报告")
    document.add_paragraph("执行摘要")
    _body(document, "本季度收入保持增长，但回款周期和重点客户续约率需要持续关注。")
    document.add_paragraph("核心结论")
    for text in ("收入同比增长 12.4%", "毛利率提升 1.8 个百分点", "应收账款周转天数增加 6 天"):
        document.add_paragraph(text, style="List Bullet")
    _landscape_section(document)
    document.add_paragraph("1 关键经营指标")
    table = document.add_table(rows=8, cols=9)
    headers = (
        "区域",
        "收入",
        "同比",
        "毛利率",
        "客户数",
        "续约率",
        "回款天数",
        "风险等级",
        "负责人",
    )
    for column, value in enumerate(headers):
        table.cell(0, column).text = value
    for row in range(1, 8):
        values = (
            f"片区{row}",
            f"{1200 + row * 137}万",
            f"{8 + row}%",
            f"{31 + row}%",
            str(90 + row * 7),
            f"{82 + row}%",
            str(38 + row),
            "中" if row % 2 else "低",
            f"经理{row}",
        )
        for column, value in enumerate(values):
            table.cell(row, column).text = value
    _fill_table_noise(table)
    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_paragraph("2 风险与行动")
    _body(document, "针对回款风险，将按客户分层设置跟进节奏并明确升级机制。")
    document.save(path)


def legal_contract(path: Path) -> None:
    document = Document()
    _add_common_noise(document, "合同文本")
    _title(document, "软件开发与技术服务合同")
    document.add_paragraph("甲方：示例科技有限公司")
    document.add_paragraph("乙方：示例软件有限公司")
    _body(document, "鉴于双方拟就软件开发与技术服务开展合作，经平等协商，达成本合同。", noisy=False)
    clauses = (
        ("第一条 合作内容", "乙方按照附件需求说明书完成系统设计、开发、测试和交付。"),
        ("第二条 项目周期", "项目自合同生效之日起九十个自然日内完成，里程碑以双方书面确认为准。"),
        ("第三条 合同价款", "合同含税总价为人民币叁拾万元整，分三期支付。"),
        ("第四条 知识产权", "项目定制成果的知识产权归甲方所有，乙方已有工具和通用组件除外。"),
        ("第五条 保密义务", "双方对履约中知悉的商业秘密承担保密义务，保密期为五年。"),
        ("第六条 违约责任", "任何一方违约均应赔偿对方因此遭受的直接损失。"),
        ("第七条 争议解决", "争议应先友好协商；协商不成，提交甲方所在地人民法院处理。"),
    )
    for heading, content in clauses:
        paragraph = document.add_paragraph(heading)
        _shade(paragraph._p, "E2F0D9")
        _body(document, content)
    table = document.add_table(rows=4, cols=4)
    for column, value in enumerate(("付款节点", "比例", "条件", "期限")):
        table.cell(0, column).text = value
    for row, values in enumerate(
        (
            ("首付款", "30%", "合同生效", "5日"),
            ("中期款", "40%", "测试通过", "10日"),
            ("尾款", "30%", "最终验收", "10日"),
        ),
        start=1,
    ):
        for column, value in enumerate(values):
            table.cell(row, column).text = value
    _fill_table_noise(table)
    document.add_paragraph("（以下无正文）")
    document.save(path)


def meeting_minutes(path: Path) -> None:
    document = Document()
    _add_common_noise(document, "会议纪要")
    _title(document, "产品迭代评审会议纪要")
    for line in (
        "会议时间：2026年8月25日 09:30-10:40",
        "会议地点：第三会议室 / 线上会议",
        "主持人：王敏",
        "参会人员：产品、设计、研发、测试代表",
    ):
        paragraph = document.add_paragraph(line)
        _shade(paragraph._p, "DDEBF7")
    document.add_paragraph("会议目标")
    _body(document, "确认默认整理模式的验收范围、风险边界与下一轮交付计划。")
    document.add_paragraph("讨论要点")
    for text in (
        "默认模式必须无需模型即可工作",
        "原文内容必须保持完整",
        "任务刷新后应能恢复",
        "错误信息应直接说明修复方法",
    ):
        document.add_paragraph(text, style="List Bullet")
    document.add_paragraph("行动项")
    table = document.add_table(rows=5, cols=5)
    for column, value in enumerate(("编号", "行动项", "负责人", "截止日期", "状态")):
        table.cell(0, column).text = value
    for row, values in enumerate(
        (
            ("A-01", "补充合同识别", "研发", "8月27日", "进行中"),
            ("A-02", "验证宽表竖版", "测试", "8月27日", "待开始"),
            ("A-03", "优化错误提示", "产品", "8月28日", "待开始"),
            ("A-04", "完成回归报告", "测试", "8月29日", "待开始"),
        ),
        start=1,
    ):
        for column, value in enumerate(values):
            table.cell(row, column).text = value
    _fill_table_noise(table)
    document.save(path)


def resume(path: Path) -> None:
    document = Document()
    _add_common_noise(document, "个人简历")
    _title(document, "陈晓 · 产品经理")
    contact = document.add_paragraph("上海｜138-0000-0000｜chenxiao@example.com｜8年经验")
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for heading, entries in (
        ("个人简介", ("专注企业效率工具与文档智能化，具备从用户研究到规模化交付的完整经验。",)),
        (
            "核心能力",
            ("产品战略与路线图", "复杂工作流设计", "跨团队项目管理", "数据分析与增长实验"),
        ),
        (
            "工作经历",
            (
                "2022.06-至今  示例科技  高级产品经理",
                "负责文档协作产品，推动活跃用户增长 45%。",
                "2018.03-2022.05  示例软件  产品经理",
                "负责企业内容管理模块和权限体系。",
            ),
        ),
        (
            "项目经历",
            ("智能文档排版平台｜负责人", "建立结构识别、规则编译、内容安全与自动验证闭环。"),
        ),
        ("教育经历", ("2014-2018  示例大学  信息管理学士",)),
    ):
        document.add_paragraph(heading)
        for entry in entries:
            _body(document, entry)
    document.save(path)


def training_manual(path: Path, image_path: Path) -> None:
    document = Document()
    _add_common_noise(document, "操作手册")
    _title(document, "客户资料导入操作手册")
    document.add_paragraph("1 使用前准备")
    _body(document, "确认账号具有管理员权限，并准备 UTF-8 编码的客户资料文件。")
    document.add_paragraph("1.1 文件要求")
    for text in ("文件格式为 CSV 或 XLSX", "首行必须为字段名称", "单次导入不超过五万行"):
        document.add_paragraph(text, style="List Bullet")
    document.add_paragraph("2 导入步骤")
    for text in (
        "进入“客户管理”页面",
        "选择“批量导入”",
        "上传文件并完成字段映射",
        "确认预检结果后开始导入",
    ):
        document.add_paragraph(text, style="List Number")
    picture = document.add_paragraph()
    picture.add_run().add_picture(str(image_path), width=Cm(6.8), height=Cm(3.2))
    document.add_paragraph("图 1 字段映射界面")
    note = document.add_paragraph(
        "注意：存在重复手机号时，系统默认更新已有记录，不会新建重复客户。"
    )
    _shade(note._p, "FCE4D6")
    document.add_paragraph("3 常见问题")
    table = document.add_table(rows=4, cols=2)
    table.cell(0, 0).text = "现象"
    table.cell(0, 1).text = "处理方法"
    issues = (
        ("上传失败", "检查文件大小和扩展名"),
        ("字段无法映射", "确认首行字段名称完整"),
        ("导入结果部分失败", "下载失败明细并修正后重试"),
    )
    for row, values in enumerate(issues, start=1):
        for column, value in enumerate(values):
            table.cell(row, column).text = value
    _fill_table_noise(table)
    document.save(path)


def financial_statement(path: Path) -> None:
    document = Document()
    _add_common_noise(document, "财务报表")
    _title(document, "2026 年度预算执行情况表")
    document.add_paragraph("编制单位：示例集团财务部   单位：万元")
    _landscape_section(document)
    table = document.add_table(rows=13, cols=11)
    headers = (
        "科目",
        "年初预算",
        "一月",
        "二月",
        "三月",
        "四月",
        "五月",
        "六月",
        "累计",
        "完成率",
        "差异说明",
    )
    for column, value in enumerate(headers):
        table.cell(0, column).text = value
    for row in range(1, 13):
        values = (
            f"成本科目{row}",
            str(800 + row * 25),
            str(60 + row),
            str(62 + row),
            str(59 + row),
            str(65 + row),
            str(68 + row),
            str(70 + row),
            str(384 + row * 6),
            f"{47 + row}%",
            "受采购节奏和项目交付周期影响",
        )
        for column, value in enumerate(values):
            table.cell(row, column).text = value
    _fill_table_noise(table)
    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_paragraph("预算执行说明")
    _body(document, "上半年整体执行进度符合计划，部分项目付款节点延后至第三季度。")
    document.save(path)


def mixed_edge_case(path: Path) -> None:
    document = Document()
    _add_common_noise(document, "综合边界")
    _title(document, "综合边界结构验证文档")
    document.add_paragraph("Overview", style="Heading 1")
    paragraph = document.add_paragraph("访问项目网站：")
    relationship_id = paragraph.part.relate_to(
        "https://example.com",
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    hyperlink_run = OxmlElement("w:r")
    hyperlink_text = OxmlElement("w:t")
    hyperlink_text.text = "https://example.com"
    hyperlink_run.append(hyperlink_text)
    hyperlink.append(hyperlink_run)
    paragraph._p.append(hyperlink)
    document.add_paragraph("1.1 中文与 English 混合标题")
    _body(document, "这是含有超链接、字段、合并单元格、嵌套表格和多节页面设置的综合样本。")
    outer = document.add_table(rows=3, cols=3)
    outer.cell(0, 0).merge(outer.cell(0, 2)).text = "合并表头"
    outer.cell(1, 0).text = "普通单元格"
    nested = outer.cell(1, 1).add_table(rows=2, cols=2)
    nested.cell(0, 0).text = "嵌套A"
    nested.cell(0, 1).text = "嵌套B"
    nested.cell(1, 0).text = "嵌套C"
    nested.cell(1, 1).text = "嵌套D"
    outer.cell(1, 2).text = "右侧内容"
    outer.cell(2, 0).text = "跨域"
    outer.cell(2, 1).text = "保护"
    outer.cell(2, 2).text = "验证"
    _fill_table_noise(outer)
    _landscape_section(document)
    document.add_paragraph("Appendix A")
    _body(document, "Landscape section should become portrait in the default cleanup mode.")
    document.save(path)


def build(output_dir: Path) -> list[Path]:
    output_dir.mkdir(parents=True, exist_ok=True)
    image_path = output_dir / "fixture-image.png"
    image_path.write_bytes(PNG_1X1)
    builders = (
        ("01-government-notice.docx", lambda path: government_notice(path)),
        ("02-academic-paper.docx", lambda path: academic_paper(path, image_path)),
        ("03-business-report.docx", lambda path: business_report(path)),
        ("04-legal-contract.docx", lambda path: legal_contract(path)),
        ("05-meeting-minutes.docx", lambda path: meeting_minutes(path)),
        ("06-resume.docx", lambda path: resume(path)),
        ("07-training-manual.docx", lambda path: training_manual(path, image_path)),
        ("08-financial-statement.docx", lambda path: financial_statement(path)),
    )
    paths: list[Path] = []
    for filename, builder in builders:
        path = output_dir / filename
        builder(path)
        paths.append(path)
    image_path.unlink()
    return paths


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--outdir", type=Path, required=True)
    args = parser.parse_args()
    for path in build(args.outdir):
        print(path)


if __name__ == "__main__":
    main()
