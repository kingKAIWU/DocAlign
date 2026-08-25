from __future__ import annotations

import base64
from pathlib import Path

import pytest
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm, Pt

PNG_1X1 = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
)


@pytest.fixture
def comprehensive_docx() -> Path:
    return Path(__file__).parent / "fixtures" / "academic-comprehensive.docx"


@pytest.fixture
def academic_docx(tmp_path: Path) -> Path:
    image_path = tmp_path / "pixel.png"
    image_path.write_bytes(PNG_1X1)
    path = tmp_path / "academic.docx"
    document = Document()

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("面向出版流程的文档格式化研究")
    title_run.bold = True
    title_run.font.size = Pt(18)

    author = document.add_paragraph("张三  示例大学")
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph("摘要")
    document.add_paragraph("本文讨论 DocAlign deterministic formatting engine。")
    document.add_paragraph("关键词：文档格式；OOXML；可复现")

    document.add_paragraph("1 Introduction", style="Heading 1")
    document.add_paragraph("1.1 Background", style="Heading 2")
    body = document.add_paragraph()
    body.add_run("正文中文与 English words 混排，")
    emphasized = body.add_run("保留强调")
    emphasized.bold = True
    body.add_run("并保持原始文本。")

    image_paragraph = document.add_paragraph()
    image_paragraph.add_run().add_picture(str(image_path), width=Mm(10))
    document.add_paragraph("图 1 DocAlign 流程")
    document.add_paragraph("表 1 核心能力")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "能力"
    table.cell(0, 1).text = "状态"
    table.cell(0, 0).merge(table.cell(0, 1))
    table.cell(1, 0).text = "内容保护"
    table.cell(1, 1).text = "通过"

    document.add_paragraph("参考文献")
    document.add_paragraph("[1] ECMA-376 Office Open XML.")
    document.sections[0].header.paragraphs[0].text = "DocAlign 测试页眉"
    document.sections[0].footer.paragraphs[0].text = "DocAlign 测试页脚"

    section = document.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    document.add_paragraph("2 Conclusion", style="Heading 1")
    document.add_paragraph("结论正文。")
    document.save(path)
    return path
