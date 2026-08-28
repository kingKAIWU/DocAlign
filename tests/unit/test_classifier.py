from __future__ import annotations

from pathlib import Path

from docalign_core.analysis.classifier import analyze_document
from docalign_core.docx.parser import parse_docx
from docalign_core.domain.document_ir import ParagraphIR
from docalign_core.domain.enums import DocumentKind, SemanticRole
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


def test_parser_preserves_block_order_and_classifier_roles(academic_docx: Path) -> None:
    parsed = parse_docx(academic_docx)
    analysis = analyze_document(parsed)
    paragraphs = [block for block in analysis.document_ir.blocks if isinstance(block, ParagraphIR)]
    roles_by_text = {paragraph.text: paragraph.detected_role for paragraph in paragraphs}
    assert roles_by_text["摘要"] == SemanticRole.ABSTRACT_HEADING
    assert roles_by_text["1 Introduction"] == SemanticRole.HEADING_1
    assert roles_by_text["1.1 Background"] == SemanticRole.HEADING_2
    assert roles_by_text["图 1 DocAlign 流程"] == SemanticRole.FIGURE_CAPTION
    assert roles_by_text["表 1 核心能力"] == SemanticRole.TABLE_CAPTION
    assert roles_by_text["参考文献"] == SemanticRole.BIBLIOGRAPHY_HEADING
    assert analysis.summary.table_count == 1
    assert analysis.summary.image_count == 1


def test_empty_paragraphs_are_not_counted_as_needing_review(tmp_path: Path) -> None:
    source = tmp_path / "empty-paragraphs.docx"
    document = Document()
    document.add_paragraph("")
    document.add_paragraph("")
    document.save(source)

    analysis = analyze_document(parse_docx(source))

    assert analysis.summary.role_counts[SemanticRole.UNKNOWN.value] == 2
    assert analysis.summary.unknown_count == 0


def test_node_ids_are_stable_for_same_source(academic_docx: Path) -> None:
    first = parse_docx(academic_docx)
    second = parse_docx(academic_docx)
    assert [block.node_id for block in first.blocks] == [block.node_id for block in second.blocks]
    assert [block.locator for block in first.blocks] == [
        block.locator for block in second.blocks
    ]
    assert [
        block.locator for block in first.blocks if block.kind == "paragraph"
    ] == [f"p{index}" for index in range(1, first.metadata.paragraph_count + 1)]
    assert [block.locator for block in first.blocks if block.kind == "table"] == ["t1"]
    paragraph = next(block for block in first.blocks if isinstance(block, ParagraphIR))
    assert [run.locator for run in paragraph.runs] == [
        f"{paragraph.locator}.r{index}" for index in range(1, len(paragraph.runs) + 1)
    ]


def test_long_numbered_sentence_is_body_not_heading(tmp_path: Path) -> None:
    source = tmp_path / "long-numbered-sentence.docx"
    text = (
        "1.1 现状分析 当前 Word 文档存在大量不规范样式：有的正文用二号字，"
        "有的标题用小五字号，中英文混排字体不区分，段落有无缩进混杂，空行杂乱无章。"
    )
    document = Document()
    document.add_paragraph(text)
    document.save(source)

    analysis = analyze_document(parse_docx(source))
    paragraph = next(
        block for block in analysis.document_ir.blocks if isinstance(block, ParagraphIR)
    )
    assert paragraph.detected_role == SemanticRole.BODY
    assert any(
        warning.code == "POSSIBLE_MIXED_HEADING_BODY" for warning in analysis.warnings
    )


def test_known_title_and_abstract_styles_update_following_context(tmp_path: Path) -> None:
    source = tmp_path / "styled-context.docx"
    document = Document()
    document.add_paragraph("智能排版研究", style="Title")
    document.add_paragraph("张三")
    document.paragraphs[-1].alignment = 1
    document.styles.add_style("DA Abstract Heading", WD_STYLE_TYPE.PARAGRAPH)
    document.add_paragraph("摘要", style="DA Abstract Heading")
    document.add_paragraph("本文研究智能文档排版方法。")
    document.add_paragraph("关键词：文档；语义；排版")
    document.add_paragraph("1 研究背景")
    document.add_paragraph("这是正文。")
    document.save(source)

    analysis = analyze_document(parse_docx(source))
    roles_by_text = {
        block.text: block.detected_role
        for block in analysis.document_ir.blocks
        if isinstance(block, ParagraphIR)
    }
    assert roles_by_text["张三"] == SemanticRole.AUTHOR_INFO
    assert roles_by_text["本文研究智能文档排版方法。"] == SemanticRole.ABSTRACT_BODY
    assert roles_by_text["这是正文。"] == SemanticRole.BODY


def test_direct_word_numbering_is_a_list_role_even_with_normal_style(tmp_path: Path) -> None:
    source = tmp_path / "direct-numbering.docx"
    document = Document()
    paragraph = document.add_paragraph("保留全部原文")
    paragraph_properties = paragraph._p.get_or_add_pPr()
    numbering = OxmlElement("w:numPr")
    level = OxmlElement("w:ilvl")
    level.set(qn("w:val"), "0")
    number_id = OxmlElement("w:numId")
    number_id.set(qn("w:val"), "1")
    numbering.extend([level, number_id])
    paragraph_properties.append(numbering)
    document.save(source)

    analysis = analyze_document(parse_docx(source))
    analyzed = next(
        block for block in analysis.document_ir.blocks if isinstance(block, ParagraphIR)
    )
    assert analyzed.current_style_name == "Normal"
    assert analyzed.detected_role == SemanticRole.LIST_ITEM
    assert analyzed.role_evidence == ["word-numbering"]


def test_relative_formatting_infers_unnumbered_heading_hierarchy(tmp_path: Path) -> None:
    source = tmp_path / "formatted-headings.docx"
    document = Document()
    body = document.add_paragraph("这是用于建立正文基准字号的普通内容。" * 8)
    body.runs[0].font.size = Pt(12)
    for text, size in (("研究框架", 18), ("数据来源", 15), ("样本筛选", 12)):
        paragraph = document.add_paragraph(text)
        paragraph.runs[0].bold = True
        paragraph.runs[0].font.size = Pt(size)
    document.save(source)

    analysis = analyze_document(parse_docx(source))
    roles_by_text = {
        block.text: block.detected_role
        for block in analysis.document_ir.blocks
        if isinstance(block, ParagraphIR)
    }
    assert roles_by_text["研究框架"] == SemanticRole.HEADING_1
    assert roles_by_text["数据来源"] == SemanticRole.HEADING_2
    assert roles_by_text["样本筛选"] == SemanticRole.HEADING_3


def test_uniform_short_paragraph_pollution_does_not_turn_resume_into_headings(
    tmp_path: Path,
) -> None:
    source = tmp_path / "polluted-resume.docx"
    document = Document()
    for text in (
        "陈晓 · 产品经理",
        "个人简介",
        "专注文档智能化",
        "核心能力",
        "产品战略",
        "跨团队协作",
        "工作经历",
        "2022.06-至今 示例科技 高级产品经理",
    ):
        paragraph = document.add_paragraph(text)
        paragraph.runs[0].bold = True
        paragraph.runs[0].font.size = Pt(15)
    document.paragraphs[0].alignment = 1
    document.save(source)

    analysis = analyze_document(parse_docx(source))
    roles_by_text = {
        block.text: block.detected_role
        for block in analysis.document_ir.blocks
        if isinstance(block, ParagraphIR)
    }
    assert analysis.summary.document_kind == DocumentKind.RESUME
    assert roles_by_text["个人简介"] == SemanticRole.HEADING_1
    assert roles_by_text["核心能力"] == SemanticRole.HEADING_1
    assert roles_by_text["专注文档智能化"] == SemanticRole.BODY
    assert roles_by_text["产品战略"] == SemanticRole.BODY


def test_contract_articles_are_left_hierarchy_and_kind_is_detected(tmp_path: Path) -> None:
    source = tmp_path / "contract.docx"
    document = Document()
    document.add_paragraph("软件服务合同")
    document.add_paragraph("甲方：示例公司")
    document.add_paragraph("乙方：服务公司")
    document.add_paragraph("第一条 服务范围")
    document.add_paragraph("乙方提供技术服务。")
    document.add_paragraph("第二条 合同价款")
    document.add_paragraph("合同价款为人民币壹万元。")
    document.save(source)

    analysis = analyze_document(parse_docx(source))
    roles_by_text = {
        block.text: block.detected_role
        for block in analysis.document_ir.blocks
        if isinstance(block, ParagraphIR)
    }
    assert analysis.summary.document_kind == DocumentKind.CONTRACT
    assert roles_by_text["第一条 服务范围"] == SemanticRole.HEADING_2
    assert roles_by_text["第二条 合同价款"] == SemanticRole.HEADING_2
