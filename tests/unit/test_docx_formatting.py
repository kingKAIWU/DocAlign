from __future__ import annotations

from docalign_core.docx.formatting import (
    apply_direct_paragraph_spec,
    clear_covered_paragraph_format,
    effective_table_font_size,
    insert_page_field,
    mm_to_twips,
    run_is_protected,
    set_repeat_table_header,
    set_row_cant_split,
    set_run_font,
    set_table_grid_borders,
    set_table_width,
    suggest_table_column_weights,
)
from docalign_core.domain.formatting_spec import (
    Alignment,
    FontSpec,
    FormattingProperty,
    LineSpacingMode,
    LineSpacingSpec,
    ParagraphSpec,
)
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def test_run_font_writes_explicit_script_fonts_and_removes_themes() -> None:
    document = Document()
    run = document.add_paragraph().add_run("中英 mixed")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:asciiTheme"), "minorHAnsi")
    fonts.set(qn("w:eastAsiaTheme"), "minorEastAsia")
    run._r.get_or_add_rPr().append(fonts)

    set_run_font(
        run,
        FontSpec(
            east_asia="宋体",
            ascii="Times New Roman",
            high_ansi="Arial",
            complex_script="Arial",
            size_pt=12,
            bold=False,
            italic=True,
            underline=True,
            color_hex="#112233",
        ),
        {
            FormattingProperty.FONT_BOLD,
            FormattingProperty.FONT_ITALIC,
            FormattingProperty.FONT_UNDERLINE,
            FormattingProperty.FONT_COLOR,
        },
    )

    written = run._r.get_or_add_rPr().find(qn("w:rFonts"))
    assert written is not None
    assert written.get(qn("w:eastAsia")) == "宋体"
    assert written.get(qn("w:ascii")) == "Times New Roman"
    assert written.get(qn("w:hAnsi")) == "Arial"
    assert written.get(qn("w:cs")) == "Arial"
    assert written.get(qn("w:asciiTheme")) is None
    assert written.get(qn("w:eastAsiaTheme")) is None
    assert run.bold is False
    assert run.italic is True
    assert run.underline is True
    assert str(run.font.color.rgb) == "112233"


def test_paragraph_setters_cover_spacing_indents_and_clear() -> None:
    document = Document()
    paragraph = document.add_paragraph("段落")
    spec = ParagraphSpec(
        alignment=Alignment.JUSTIFY,
        line_spacing=LineSpacingSpec(mode=LineSpacingMode.EXACT, value=18),
        space_before_pt=6,
        space_after_pt=3,
        hanging_indent_pt=12,
        left_indent_pt=9,
        right_indent_pt=4,
        keep_with_next=True,
        keep_lines_together=True,
        page_break_before=True,
        widow_orphan_control=False,
    )
    apply_direct_paragraph_spec(paragraph, spec)

    assert paragraph.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY
    assert paragraph.paragraph_format.line_spacing_rule == WD_LINE_SPACING.EXACTLY
    assert paragraph.paragraph_format.first_line_indent is not None
    assert paragraph.paragraph_format.first_line_indent.pt == -12
    assert paragraph.paragraph_format.keep_with_next is True
    widow = paragraph._p.get_or_add_pPr().find(qn("w:widowControl"))
    assert widow is not None and widow.get(qn("w:val")) == "0"

    clear_covered_paragraph_format(paragraph, spec)
    assert paragraph.alignment is None
    assert paragraph.paragraph_format.space_before is None
    assert paragraph.paragraph_format.first_line_indent is None
    assert paragraph._p.get_or_add_pPr().find(qn("w:widowControl")) is None


def test_multiple_and_at_least_line_spacing_are_applied() -> None:
    document = Document()
    multiple = document.add_paragraph("multiple")
    apply_direct_paragraph_spec(
        multiple,
        ParagraphSpec(line_spacing=LineSpacingSpec(mode=LineSpacingMode.MULTIPLE, value=1.5)),
    )
    assert multiple.paragraph_format.line_spacing_rule == WD_LINE_SPACING.ONE_POINT_FIVE

    at_least = document.add_paragraph("at least")
    apply_direct_paragraph_spec(
        at_least,
        ParagraphSpec(line_spacing=LineSpacingSpec(mode=LineSpacingMode.AT_LEAST, value=14)),
    )
    assert at_least.paragraph_format.line_spacing_rule == WD_LINE_SPACING.AT_LEAST


def test_table_width_and_page_field_are_deterministic() -> None:
    document = Document()
    table = document.add_table(rows=1, cols=1)
    set_table_width(table, mm_to_twips(100))
    width = table._tbl.tblPr.find(qn("w:tblW"))
    assert width is not None
    assert width.get(qn("w:type")) == "dxa"
    assert int(width.get(qn("w:w"), "0")) > 0
    assert table.autofit is False

    paragraph = document.add_paragraph()
    assert insert_page_field(paragraph) is True
    assert insert_page_field(paragraph) is False
    assert run_is_protected(paragraph.runs[0]) is True
    instructions = [node.text for node in paragraph._p.iter(qn("w:instrText"))]
    assert instructions == [" PAGE "]


def test_table_cleanup_helpers_create_readable_cross_page_geometry() -> None:
    document = Document()
    table = document.add_table(rows=2, cols=3)
    table.cell(0, 0).text = "科目"
    table.cell(0, 1).text = "数值"
    table.cell(0, 2).text = "差异说明"
    table.cell(1, 0).text = "成本科目"
    table.cell(1, 1).text = "100"
    table.cell(1, 2).text = "受采购节奏和项目交付周期影响"

    weights = suggest_table_column_weights(table)
    assert weights[2] > weights[1]
    set_table_width(
        table,
        mm_to_twips(160),
        column_weights=weights,
        min_column_width_twips=mm_to_twips(8),
    )
    set_table_grid_borders(table)
    set_repeat_table_header(table.rows[0])
    for row in table.rows:
        set_row_cant_split(row)

    assert table._tbl.tblPr.find(qn("w:tblBorders")) is not None
    assert table.rows[0]._tr.trPr.find(qn("w:tblHeader")) is not None
    assert all(row._tr.trPr.find(qn("w:cantSplit")) is not None for row in table.rows)
    assert effective_table_font_size(10.5, 11, adaptive=True, minimum_size_pt=8) == 8
    assert effective_table_font_size(10.5, 4, adaptive=True, minimum_size_pt=8) == 10.5
