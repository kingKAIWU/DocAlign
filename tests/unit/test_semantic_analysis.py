from __future__ import annotations

import json
from pathlib import Path

import httpx
import pytest
from docalign_core.analysis.classifier import analyze_document
from docalign_core.analysis.semantic import (
    SemanticAnalysisDraft,
    SemanticRoleAssignment,
    merge_semantic_analysis,
)
from docalign_core.docx.parser import parse_docx
from docalign_core.domain.document_ir import AnalysisResult, ParagraphIR
from docalign_core.domain.enums import AnalysisMode, DocumentKind, RoleSource, SemanticRole
from docalign_core.llm.semantic import OpenAICompatibleSemanticAnalyzer
from docx import Document


def _plain_document(tmp_path: Path) -> tuple[Path, AnalysisResult]:
    source = tmp_path / "plain.docx"
    document = Document()
    document.add_paragraph("智能文档排版研究")
    document.add_paragraph("研究背景")
    document.add_paragraph("这是介绍研究背景的普通正文段落。")
    document.add_paragraph("1 已有编号标题")
    document.save(source)
    return source, analyze_document(parse_docx(source))


def test_semantic_merge_upgrades_ambiguous_plain_text_but_locks_numbering(
    tmp_path: Path,
) -> None:
    _, deterministic = _plain_document(tmp_path)
    paragraphs = [
        block
        for block in deterministic.document_ir.blocks
        if isinstance(block, ParagraphIR)
    ]
    draft = SemanticAnalysisDraft(
        document_kind=DocumentKind.ACADEMIC_PAPER,
        document_kind_confidence=0.91,
        assignments=[
            SemanticRoleAssignment(
                node_id=paragraphs[0].node_id,
                role=SemanticRole.TITLE,
                confidence=0.94,
                evidence="opening document title",
            ),
            SemanticRoleAssignment(
                node_id=paragraphs[1].node_id,
                role=SemanticRole.HEADING_1,
                confidence=0.9,
                evidence="short section label followed by prose",
            ),
            SemanticRoleAssignment(
                node_id=paragraphs[2].node_id,
                role=SemanticRole.BODY,
                confidence=0.96,
                evidence="complete explanatory sentence",
            ),
            SemanticRoleAssignment(
                node_id=paragraphs[3].node_id,
                role=SemanticRole.BODY,
                confidence=0.99,
                evidence="model mistake",
            ),
        ],
    )

    merged = merge_semantic_analysis(
        deterministic, draft, provider="mock", model="semantic-test"
    )
    merged_paragraphs = [
        block for block in merged.document_ir.blocks if isinstance(block, ParagraphIR)
    ]
    assert merged_paragraphs[0].detected_role == SemanticRole.TITLE
    assert merged_paragraphs[1].detected_role == SemanticRole.HEADING_1
    assert merged_paragraphs[1].role_source == RoleSource.LLM
    assert merged_paragraphs[3].detected_role == SemanticRole.HEADING_1
    assert any(w.code == "SEMANTIC_OVERRIDE_REJECTED" for w in merged.warnings)
    assert merged.summary.analysis_mode == AnalysisMode.SMART
    assert merged.summary.document_kind == DocumentKind.ACADEMIC_PAPER
    assert merged.summary.model_reviewed_paragraphs == 4


@pytest.mark.asyncio
async def test_deepseek_semantic_analyzer_sends_paragraphs_and_uses_json_object(
    tmp_path: Path,
) -> None:
    _, deterministic = _plain_document(tmp_path)
    seen: dict[str, object] = {}
    paragraph = next(
        block
        for block in deterministic.document_ir.blocks
        if isinstance(block, ParagraphIR)
    )

    def handler(request: httpx.Request) -> httpx.Response:
        seen.update(json.loads(request.content))
        content = SemanticAnalysisDraft(
            document_kind=DocumentKind.REPORT,
            document_kind_confidence=0.8,
            assignments=[
                SemanticRoleAssignment(
                    node_id=paragraph.node_id,
                    role=SemanticRole.TITLE,
                    confidence=0.9,
                    evidence="opening title",
                )
            ],
        ).model_dump_json()
        return httpx.Response(200, json={"choices": [{"message": {"content": content}}]})

    async with httpx.AsyncClient(transport=httpx.MockTransport(handler)) as client:
        analyzer = OpenAICompatibleSemanticAnalyzer(
            base_url="https://api.deepseek.com/v1",
            model="deepseek-v4-flash",
            client=client,
        )
        draft = await analyzer.analyze(deterministic.document_ir, deterministic)

    assert draft.document_kind == DocumentKind.REPORT
    assert seen["response_format"] == {"type": "json_object"}
    assert seen["thinking"] == {"type": "disabled"}
    messages = seen["messages"]
    assert isinstance(messages, list)
    user_payload = json.loads(messages[-1]["content"])
    assert user_payload["paragraphs"][0]["text"] == "智能文档排版研究"
    assert "runs" not in user_payload["paragraphs"][0]
