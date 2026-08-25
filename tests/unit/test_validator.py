from __future__ import annotations

from pathlib import Path

from docalign_core.analysis.classifier import analyze_document
from docalign_core.docx.parser import parse_docx
from docalign_core.domain.enums import Severity
from docalign_core.domain.formatting_spec import default_cleanup_spec
from docalign_core.services.processing import process_document
from docalign_core.validation.validator import DocumentValidator
from docx import Document
from docx.oxml.ns import qn


def test_validator_finds_table_font_and_header_errors_with_stable_locators(
    academic_docx: Path,
    tmp_path: Path,
) -> None:
    analysis = analyze_document(parse_docx(academic_docx))
    spec = default_cleanup_spec()
    output = tmp_path / "formatted.docx"
    result = process_document(
        academic_docx,
        analysis.document_ir,
        spec,
        output,
        job_id="job-table-validation",
        artifact_dir=tmp_path / "artifacts",
    )
    assert result.audit.validation.valid

    corrupted = Document(output)
    table = corrupted.tables[0]
    header = table.rows[0]._tr.trPr.find(qn("w:tblHeader"))
    assert header is not None
    table.rows[0]._tr.trPr.remove(header)
    run = next(
        run
        for row in table.rows
        for cell in row.cells
        for paragraph in cell.paragraphs
        for run in paragraph.runs
        if run.text
    )
    fonts = run._r.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:eastAsia"), "黑体")
    corrupted.save(output)

    validation = DocumentValidator().validate(output, spec, analysis.document_ir)

    assert not validation.valid
    assert any(
        issue.code == "TABLE_HEADER_REPEAT_MISSING" and issue.locator == "t1.r1"
        for issue in validation.issues
    )
    assert any(
        issue.code == "TABLE_FONT_VALIDATION_FAILED"
        and issue.locator is not None
        and issue.locator.startswith("t1.r1.c1.p1.r1")
        for issue in validation.issues
    )
    assert all(
        issue.severity in {Severity.ERROR, Severity.FATAL}
        for issue in validation.issues
    )
