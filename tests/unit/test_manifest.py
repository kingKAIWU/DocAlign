from __future__ import annotations

from pathlib import Path

from docalign_core.docx.manifest import extract_format_manifest
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Mm, Pt


def test_format_manifest_is_deterministic_and_evidence_backed(tmp_path: Path) -> None:
    source = tmp_path / "reference-format.docx"
    document = Document()
    section = document.sections[0]
    section.left_margin = Mm(28)
    section.right_margin = Mm(24)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    fonts = normal.element.get_or_add_rPr().rFonts
    assert fonts is not None
    fonts.set(qn("w:eastAsia"), "宋体")

    document.add_paragraph("格式画像示例", style="Title")
    document.add_paragraph("这是用于提取样式证据的正文。")
    table = document.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.cell(0, 0).text = "项目"
    table.cell(0, 1).text = "要求"
    table.cell(1, 0).text = "字体"
    table.cell(1, 1).text = "宋体"
    header = section.header.paragraphs[0]
    header.text = "内部资料"
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.save(source)

    first = extract_format_manifest(
        source,
        document_id="doc_manifest",
        source_filename="参考模板.docx",
    )
    second = extract_format_manifest(
        source,
        document_id="doc_manifest",
        source_filename="参考模板.docx",
    )

    assert first == second
    assert first.source_filename == "参考模板.docx"
    assert first.summary.requirement_count == len(first.requirements)
    assert first.summary.by_category["style"] > 0
    assert first.summary.by_category["section"] > 0
    assert first.summary.by_category["table"] > 0
    assert first.summary.by_category["header_footer"] > 0
    assert [item.requirement_id for item in first.requirements] == [
        f"R{index:04d}" for index in range(1, len(first.requirements) + 1)
    ]
    assert all(item.source_part and item.evidence for item in first.requirements)
    assert any(
        item.target == "style:Normal"
        and item.property_path == "font.east_asia"
        and item.expected == "宋体"
        and item.auto_applicable
        for item in first.requirements
    )
    assert any(
        item.target == "t1"
        and item.property_path == "table.alignment"
        and item.expected == "center"
        and not item.auto_applicable
        for item in first.requirements
    )
