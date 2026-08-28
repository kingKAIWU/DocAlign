from __future__ import annotations

from pathlib import Path

from docalign_core.docx.template_candidate import compile_template_rule_candidate
from docalign_core.domain.enums import SemanticRole
from docalign_core.domain.formatting_spec import Orientation, PageSize, SpecSourceType
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Mm, Pt


def test_template_candidate_compiles_only_safe_role_and_page_rules(tmp_path: Path) -> None:
    source = tmp_path / "approved-reference.docx"
    document = Document()
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(26)
    section.bottom_margin = Mm(24)
    section.left_margin = Mm(30)
    section.right_margin = Mm(22)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal_fonts = normal.element.get_or_add_rPr().rFonts
    assert normal_fonts is not None
    normal_fonts.set(qn("w:eastAsia"), "宋体")
    normal.paragraph_format.first_line_indent = Pt(24)
    normal.paragraph_format.line_spacing = 1.5

    title = document.styles["Title"]
    title.font.size = Pt(18)
    title.font.bold = True
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    heading = document.styles["Heading 1"]
    heading.font.size = Pt(16)
    heading.font.bold = True
    heading.paragraph_format.keep_with_next = True

    document.add_paragraph("经确认的报告样例", style="Title")
    document.add_paragraph("一、研究背景", style="Heading 1")
    document.add_paragraph("这是第一段正文，使用参考文档的正文样式。")
    document.add_paragraph("这是第二段正文，用于确认样式确实在文档中使用。")
    document.add_table(rows=1, cols=2)
    document.save(source)

    candidate = compile_template_rule_candidate(source, source_filename="合格样例.docx")

    assert candidate.safe_to_apply
    assert candidate.source_filename == "合格样例.docx"
    assert candidate.spec.source.type == SpecSourceType.TEMPLATE
    assert candidate.spec.source.reference_sha256 == candidate.source_sha256
    assert candidate.spec.auto_layout.enabled is False
    assert candidate.spec.document is not None
    assert candidate.spec.document.page.size == PageSize.A4
    assert candidate.spec.document.page.orientation == Orientation.PORTRAIT
    assert candidate.spec.document.page.margin_left_mm == 30
    assert candidate.spec.roles[SemanticRole.BODY].font is not None
    assert candidate.spec.roles[SemanticRole.BODY].font.east_asia == "宋体"
    assert candidate.spec.roles[SemanticRole.BODY].paragraph is not None
    assert candidate.spec.roles[SemanticRole.BODY].paragraph.first_line_indent_pt == 24
    assert candidate.spec.roles[SemanticRole.HEADING_1].font is not None
    assert candidate.spec.roles[SemanticRole.HEADING_1].font.size_pt == 16
    assert candidate.spec.roles[SemanticRole.TITLE].paragraph is not None
    assert candidate.spec.roles[SemanticRole.TITLE].paragraph.alignment == "center"
    assert {mapping.role for mapping in candidate.role_mappings} >= {
        SemanticRole.BODY,
        SemanticRole.TITLE,
        SemanticRole.HEADING_1,
    }
    assert any("表格属性" in item for item in candidate.unsupported_features)
    assert candidate.summary.applied_requirement_count == len(candidate.applied_requirement_ids)
    assert 0 < candidate.summary.coverage_percent <= 100


def test_template_candidate_does_not_guess_unmapped_custom_body_style(tmp_path: Path) -> None:
    source = tmp_path / "ambiguous-reference.docx"
    document = Document()
    custom = document.styles.add_style("内部专用", WD_STYLE_TYPE.PARAGRAPH)
    custom.font.size = Pt(13)
    document.add_paragraph("普通内容，没有足够证据说明它对应哪个规则角色。", style=custom)
    document.save(source)

    candidate = compile_template_rule_candidate(source, source_filename="待判断样例.docx")

    assert SemanticRole.BODY not in candidate.spec.roles
    assert any("内部专用" in item for item in candidate.ambiguities)
