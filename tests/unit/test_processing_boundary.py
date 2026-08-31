from __future__ import annotations

from pathlib import Path

from docalign_core.analysis.processing_boundary import build_processing_boundary
from docalign_core.docx.parser import parse_docx
from docalign_core.domain.document_ir import (
    DocumentFeatureHandling,
    PackagePartIR,
    RelationshipIR,
    UnsupportedBlockIR,
)
from docx import Document
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement


def test_processing_boundary_exposes_complex_content_and_required_review(
    tmp_path: Path,
) -> None:
    source = tmp_path / "complex-boundary.docx"
    document = Document()
    field = document.add_paragraph("目录")
    field._p.append(OxmlElement("w:fldSimple"))
    equation = document.add_paragraph("公式")
    equation._p.append(OxmlElement("m:oMath"))
    text_box = document.add_paragraph("文本框")
    text_box._p.append(OxmlElement("w:txbxContent"))
    content_control = document.add_paragraph("受控内容")
    content_control._p.append(OxmlElement("w:sdt"))
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).merge(table.cell(0, 1))
    table.cell(1, 0).add_table(rows=1, cols=1)
    document.add_section(WD_SECTION.NEW_PAGE)
    document.save(source)

    document_ir = parse_docx(source)
    assert document_ir.metadata.feature_counts["field"] == 1
    assert document_ir.metadata.feature_counts["equation"] == 1
    assert document_ir.metadata.feature_counts["text_box"] == 1
    assert document_ir.metadata.feature_counts["content_control"] == 1

    document_ir.package_parts.extend(
        [
            _part("word/footnotes.xml"),
            _part("word/comments.xml"),
            _part("word/embeddings/object1.bin"),
            _part("customXml/item1.xml"),
        ]
    )
    document_ir.relationships.append(
        RelationshipIR(
            source_part="/word/document.xml",
            relationship_id="rIdExternal",
            relationship_type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            target="https://example.com",
            external=True,
        )
    )
    document_ir.blocks.append(
        UnsupportedBlockIR(
            node_id="unsupported-1",
            locator="u1",
            index=len(document_ir.blocks),
            qname="{urn:test}unknown",
            text_preview="受保护结构",
        )
    )

    boundary = build_processing_boundary(document_ir)
    items = {item.code: item for item in boundary.items}

    assert boundary.status == "review_required"
    assert boundary.acknowledgment_required is True
    assert boundary.review_feature_count >= 9
    assert items["field"].locators == ["p1"]
    assert items["merged_table"].locators == ["t1"]
    assert items["nested_table"].locators == ["t1"]
    assert items["unknown_ooxml"].locators == ["u1"]
    assert items["footnote"].handling == DocumentFeatureHandling.PRESERVE_ONLY
    assert items["embedded_object"].review_required is True
    assert items["multiple_sections"].count == 2


def test_processing_boundary_keeps_simple_document_friction_free(tmp_path: Path) -> None:
    source = tmp_path / "simple.docx"
    document = Document()
    document.add_paragraph("普通正文")
    document.save(source)

    boundary = build_processing_boundary(parse_docx(source))

    assert boundary.status == "standard"
    assert boundary.detected_feature_count == 0
    assert boundary.review_feature_count == 0
    assert boundary.acknowledgment_required is False


def _part(path: str) -> PackagePartIR:
    return PackagePartIR(
        path=path,
        compressed_size=1,
        uncompressed_size=1,
        sha256="a" * 64,
    )
