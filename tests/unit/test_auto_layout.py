from __future__ import annotations

from pathlib import Path

import pytest
from docalign_core.analysis.classifier import analyze_document
from docalign_core.docx.parser import parse_docx
from docalign_core.domain.enums import SemanticRole
from docalign_core.domain.formatting_spec import (
    AutoLayoutSpec,
    default_academic_spec,
    default_cleanup_spec,
)
from docalign_core.engine.auto_layout import (
    AutoLayoutIntegrityError,
    apply_auto_layout,
    assert_auto_layout_integrity,
)
from docalign_core.services.processing import process_document
from docx import Document
from docx.oxml import OxmlElement


def test_auto_layout_splits_manual_lines_and_reclassifies_heading_levels(
    tmp_path: Path,
) -> None:
    source = tmp_path / "continuous.docx"
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("1 研究背景")
    paragraph.add_run().add_break()
    paragraph.add_run("这是第一段正文。")
    paragraph.add_run().add_break()
    paragraph.add_run("1.1 技术方法")
    paragraph.add_run().add_break()
    paragraph.add_run("这是第二段正文。")
    document.save(source)

    source_ir = analyze_document(parse_docx(source)).document_ir
    structured_path = tmp_path / "structured.docx"
    result = apply_auto_layout(
        source,
        source_ir,
        AutoLayoutSpec(enabled=True),
        structured_path,
    )

    assert len(result.changes) == 1
    assert result.document_ir.metadata.paragraph_count == 4
    paragraphs = [block for block in result.document_ir.blocks if block.kind == "paragraph"]
    assert [block.detected_role for block in paragraphs] == [
        SemanticRole.HEADING_1,
        SemanticRole.BODY,
        SemanticRole.HEADING_2,
        SemanticRole.BODY,
    ]
    assert "".join(result.document_ir.content_fingerprint.paragraph_texts) == (
        "1 研究背景这是第一段正文。1.1 技术方法这是第二段正文。"
    )


def test_auto_layout_splits_long_body_at_sentence_boundaries(tmp_path: Path) -> None:
    source = tmp_path / "long-body.docx"
    sentence = "这一句用于验证正文会在完整句子边界处分段，并且不会改写任何字符。"
    original_text = sentence * 8
    document = Document()
    document.add_paragraph(original_text)
    document.save(source)

    source_ir = analyze_document(parse_docx(source)).document_ir
    result = apply_auto_layout(
        source,
        source_ir,
        AutoLayoutSpec(enabled=True, target_body_chars=80, max_body_chars=120),
        tmp_path / "long-body-structured.docx",
    )

    assert len(result.changes) == 1
    texts = result.document_ir.content_fingerprint.paragraph_texts
    assert len(texts) > 1
    assert "".join(texts) == original_text
    assert all(text.endswith("。") for text in texts)


def test_auto_layout_recovers_a_collapsed_numbered_outline_with_render_markers(
    tmp_path: Path,
) -> None:
    source = tmp_path / "collapsed-outline.docx"
    original_text = (
        "智能图像研究报告"
        "一、绪论"
        "1.1 研究背景随着技术快速发展，图像处理已经成为重要研究方向。"
        "1.2 研究现状目前相关方法仍有进一步提升空间。"
        "二、技术方法"
        "2.1 技术路线本文采用空频双分支结构完成特征建模。"
        "2.2 实施方案本研究通过分阶段实验验证方法有效性。"
    )
    document = Document()
    paragraph = document.add_paragraph()
    run = paragraph.add_run(original_text)
    run._r.append(OxmlElement("w:lastRenderedPageBreak"))
    document.save(source)

    source_ir = analyze_document(parse_docx(source)).document_ir
    result = apply_auto_layout(
        source,
        source_ir,
        AutoLayoutSpec(enabled=True),
        tmp_path / "collapsed-outline-structured.docx",
    )

    assert len(result.changes) == 1
    assert result.changes[0].reason == "collapsed numbered outline"
    paragraphs = [block for block in result.document_ir.blocks if block.kind == "paragraph"]
    assert [block.detected_role for block in paragraphs] == [
        SemanticRole.TITLE,
        SemanticRole.HEADING_1,
        SemanticRole.HEADING_2,
        SemanticRole.BODY,
        SemanticRole.HEADING_2,
        SemanticRole.BODY,
        SemanticRole.HEADING_1,
        SemanticRole.HEADING_2,
        SemanticRole.BODY,
        SemanticRole.HEADING_2,
        SemanticRole.BODY,
    ]
    assert "".join(result.document_ir.content_fingerprint.paragraph_texts) == original_text
    assert not any(
        warning.code == "AUTO_LAYOUT_PROTECTED_PARAGRAPH_SKIPPED"
        for warning in result.warnings
    )

    formatted = tmp_path / "collapsed-outline-formatted.docx"
    processed = process_document(
        source,
        source_ir,
        default_cleanup_spec(),
        formatted,
        job_id="job-collapsed-outline",
        artifact_dir=tmp_path / "collapsed-outline-artifacts",
    )
    assert processed.audit.validation.valid
    assert processed.audit.summary.paragraphs_before == 1
    assert processed.audit.summary.paragraphs_after == 11
    assert [paragraph.style.name for paragraph in Document(formatted).paragraphs] == [
        "DA Title",
        "DA Heading 1",
        "DA Heading 2",
        "DA Body",
        "DA Heading 2",
        "DA Body",
        "DA Heading 1",
        "DA Heading 2",
        "DA Body",
        "DA Heading 2",
        "DA Body",
    ]


def test_auto_layout_skips_paragraphs_with_protected_fields(tmp_path: Path) -> None:
    source = tmp_path / "protected.docx"
    document = Document()
    paragraph = document.add_paragraph("受保护字段前的长正文。" * 30)
    field_run = paragraph.add_run()
    field_run._r.append(OxmlElement("w:fldChar"))
    document.save(source)

    source_ir = analyze_document(parse_docx(source)).document_ir
    result = apply_auto_layout(
        source,
        source_ir,
        AutoLayoutSpec(enabled=True, target_body_chars=80, max_body_chars=120),
        tmp_path / "protected-structured.docx",
    )

    assert not result.changes
    assert result.source_path == source
    assert any(
        warning.code == "AUTO_LAYOUT_PROTECTED_PARAGRAPH_SKIPPED"
        for warning in result.warnings
    )


def test_processing_audits_structural_layout_and_formats_new_paragraphs(
    tmp_path: Path,
) -> None:
    source = tmp_path / "workflow.docx"
    document = Document()
    paragraph = document.add_paragraph("1 总览")
    paragraph.add_run().add_break()
    paragraph.add_run("正文第一段。")
    paragraph.add_run().add_break()
    paragraph.add_run("1.1 细节")
    paragraph.add_run().add_break()
    paragraph.add_run("正文第二段。")
    document.save(source)

    analysis = analyze_document(parse_docx(source))
    output = tmp_path / "formatted.docx"
    result = process_document(
        source,
        analysis.document_ir,
        default_academic_spec(),
        output,
        job_id="job-auto-layout",
        artifact_dir=tmp_path / "artifacts",
    )

    assert result.audit.validation.valid
    assert result.auto_layout_splits == 1
    assert result.audit.summary.paragraphs_before == 1
    assert result.audit.summary.paragraphs_after == 4
    assert [paragraph.style.name for paragraph in Document(output).paragraphs] == [
        "DA Heading 1",
        "DA Body",
        "DA Heading 2",
        "DA Body",
    ]


def test_integrity_guard_rejects_text_changes(tmp_path: Path) -> None:
    before_path = tmp_path / "before.docx"
    after_path = tmp_path / "after.docx"
    before = Document()
    before.add_paragraph("原始正文")
    before.save(before_path)
    after = Document()
    after.add_paragraph("被改写的正文")
    after.save(after_path)

    with pytest.raises(AutoLayoutIntegrityError, match="main_text"):
        assert_auto_layout_integrity(parse_docx(before_path), parse_docx(after_path))
