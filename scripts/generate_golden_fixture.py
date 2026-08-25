from __future__ import annotations

import argparse
import base64
from datetime import UTC, datetime
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt

PNG_1X1 = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
)


def build_fixture(target: Path) -> None:
    target.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    fixed_time = datetime(2020, 1, 1, tzinfo=UTC)
    document.core_properties.author = "DocAlign"
    document.core_properties.created = fixed_time
    document.core_properties.modified = fixed_time
    document.settings.odd_and_even_pages_header_footer = True
    _configure_cjk_fonts(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("面向出版流程的文档格式化研究")
    title_run.bold = True
    title_run.font.size = Pt(18)
    document.add_paragraph("张三  示例大学").alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph("摘要")
    document.add_paragraph("本文讨论 DocAlign deterministic formatting engine。")
    document.add_paragraph("关键词：文档格式；OOXML；可复现")
    document.add_paragraph("1 Introduction", style="Heading 1")
    document.add_paragraph("1.1 Background", style="Heading 2")

    body = document.add_paragraph()
    body.add_run("正文中文与 English words 混排，")
    emphasized = body.add_run("保留强调")
    emphasized.bold = True
    body.add_run("、书签、")
    bookmarked = body.add_run("锚点内容")
    _wrap_bookmark(bookmarked, "research_anchor", "42")
    body.add_run("与原始文本。")

    link_paragraph = document.add_paragraph("外部资料：")
    _append_hyperlink(link_paragraph, "ECMA-376", "https://ecma-international.org/")

    field_paragraph = document.add_paragraph("图序号：")
    _append_field(field_paragraph, " SEQ Figure \\* ARABIC ", "1")

    equation = document.add_paragraph("公式：")
    math = OxmlElement("m:oMath")
    math_run = OxmlElement("m:r")
    math_text = OxmlElement("m:t")
    math_text.text = "E=mc²"
    math_run.append(math_text)
    math.append(math_run)
    equation._p.append(math)

    image_paragraph = document.add_paragraph()
    image_paragraph.add_run().add_picture(BytesIO(PNG_1X1), width=Mm(10))
    document.add_paragraph("图 1 DocAlign 流程")
    document.add_paragraph("表 1 核心能力")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "能力"
    table.cell(0, 1).text = "状态"
    table.cell(0, 0).merge(table.cell(0, 1))
    table.cell(1, 0).text = "内容保护"
    table.cell(1, 1).text = "通过"
    nested = table.cell(1, 0).add_table(rows=1, cols=1)
    nested.cell(0, 0).text = "嵌套说明"

    document.add_paragraph("参考文献")
    document.add_paragraph("[1] ECMA-376 Office Open XML.")
    first_section = document.sections[0]
    first_section.different_first_page_header_footer = True
    first_section.header.paragraphs[0].text = "DocAlign 测试页眉"
    first_section.first_page_header.paragraphs[0].text = "DocAlign 首页页眉"
    first_section.even_page_header.paragraphs[0].text = "DocAlign 偶数页页眉"
    footer = first_section.footer.paragraphs[0]
    footer.add_run("第 ")
    _append_field(footer, " PAGE ", "1")
    footer.add_run(" 页")

    section = document.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    document.add_paragraph("2 Conclusion", style="Heading 1")
    document.add_paragraph("结论正文。")
    _append_content_control(document)

    document.save(target)


def _append_hyperlink(paragraph: object, text: str, url: str) -> None:
    relationship_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.extend([color, underline])
    value = OxmlElement("w:t")
    value.text = text
    run.extend([properties, value])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _configure_cjk_fonts(document: object) -> None:
    for style in document.styles:
        if style.type != WD_STYLE_TYPE.PARAGRAPH:
            continue
        style.font.name = "Times New Roman"
        properties = style.element.get_or_add_rPr()
        fonts = properties.find(qn("w:rFonts"))
        if fonts is None:
            fonts = OxmlElement("w:rFonts")
            properties.insert(0, fonts)
        fonts.set(qn("w:ascii"), "Times New Roman")
        fonts.set(qn("w:hAnsi"), "Times New Roman")
        fonts.set(qn("w:eastAsia"), "Arial Unicode MS")


def _append_field(paragraph: object, instruction_text: str, display_text: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instruction.text = instruction_text
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = display_text
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, value, end])


def _wrap_bookmark(run: object, name: str, bookmark_id: str) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), bookmark_id)
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), bookmark_id)
    run._r.addprevious(start)
    run._r.addnext(end)


def _append_content_control(document: object) -> None:
    control = OxmlElement("w:sdt")
    properties = OxmlElement("w:sdtPr")
    tag = OxmlElement("w:tag")
    tag.set(qn("w:val"), "DocAlignFixture")
    properties.append(tag)
    content = OxmlElement("w:sdtContent")
    paragraph = OxmlElement("w:p")
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "受保护的未知顶层 OOXML 内容"
    run.append(text)
    paragraph.append(run)
    content.append(paragraph)
    control.extend([properties, content])
    document._element.body.insert(-1, control)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--out",
        type=Path,
        default=Path("tests/fixtures/academic-comprehensive.docx"),
    )
    args = parser.parse_args()
    build_fixture(args.out)


if __name__ == "__main__":
    main()
