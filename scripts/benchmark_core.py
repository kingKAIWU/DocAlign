from __future__ import annotations

import argparse
import json
import tempfile
from pathlib import Path
from time import perf_counter

from docalign_core.analysis.classifier import analyze_document
from docalign_core.docx.parser import parse_docx
from docalign_core.domain.formatting_spec import default_academic_spec
from docalign_core.services.processing import process_document
from docx import Document


def _elapsed(started_at: float) -> float:
    return round(perf_counter() - started_at, 3)


def _build_document(path: Path, paragraph_count: int) -> float:
    started_at = perf_counter()
    document = Document()
    document.add_heading("DocAlign 性能基准", level=1)
    for index in range(paragraph_count - 1):
        document.add_paragraph(
            f"第 {index + 1} 段：中英文 mixed content for deterministic formatting benchmark."
        )
    document.save(str(path))
    return _elapsed(started_at)


def run_benchmark(paragraph_count: int, max_core_seconds: float) -> dict[str, object]:
    with tempfile.TemporaryDirectory(prefix="docalign-benchmark-") as directory:
        root = Path(directory)
        source = root / "source.docx"
        output = root / "output.docx"
        generate_seconds = _build_document(source, paragraph_count)

        started_at = perf_counter()
        document_ir = parse_docx(source, document_id="benchmark-document")
        parse_seconds = _elapsed(started_at)

        started_at = perf_counter()
        analysis = analyze_document(document_ir)
        analyze_seconds = _elapsed(started_at)

        started_at = perf_counter()
        result = process_document(
            source,
            analysis.document_ir,
            default_academic_spec(),
            output,
            job_id="benchmark-job",
            artifact_dir=root / "artifacts",
        )
        format_validate_seconds = _elapsed(started_at)
        core_seconds = round(parse_seconds + analyze_seconds + format_validate_seconds, 3)

        payload: dict[str, object] = {
            "paragraphs": analysis.summary.paragraph_count,
            "source_size_bytes": source.stat().st_size,
            "generate_seconds": generate_seconds,
            "parse_seconds": parse_seconds,
            "analyze_seconds": analyze_seconds,
            "format_validate_seconds": format_validate_seconds,
            "core_seconds": core_seconds,
            "validation_valid": result.audit.validation.valid,
            "mutations": len(result.audit.mutations),
            "max_core_seconds": max_core_seconds,
            "within_target": core_seconds <= max_core_seconds,
        }
        if not result.audit.validation.valid:
            raise SystemExit("Benchmark output failed validation.")
        if core_seconds > max_core_seconds:
            raise SystemExit(json.dumps(payload, ensure_ascii=False, indent=2))
        return payload


def main() -> None:
    parser = argparse.ArgumentParser(description="Benchmark the deterministic DocAlign core.")
    parser.add_argument("--paragraphs", type=int, default=2_000)
    parser.add_argument("--max-core-seconds", type=float, default=30.0)
    args = parser.parse_args()
    if args.paragraphs < 1:
        parser.error("--paragraphs must be at least 1")
    if args.max_core_seconds <= 0:
        parser.error("--max-core-seconds must be positive")
    print(
        json.dumps(
            run_benchmark(args.paragraphs, args.max_core_seconds),
            ensure_ascii=False,
            indent=2,
        )
    )


if __name__ == "__main__":
    main()
