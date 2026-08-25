from __future__ import annotations

import argparse
import base64
from datetime import UTC, datetime
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt
from docx.text.paragraph import Paragraph

PNG_1X1 = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
)
FIXED_TIME = datetime(2020, 1, 1, tzinfo=UTC)


def _document() -> DocumentObject:
    document = Document()
    document.core_properties.author = "DocAlign"
    document.core_properties.created = FIXED_TIME
    document.core_properties.modified = FIXED_TIME
    return document


def _save(document: DocumentObject, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(path))


def _mixed_body(path: Path) -> None:
    document = _document()
    section = document.sections[0]
    section.left_margin = Mm(33)
    section.right_margin = Mm(27)
    title = document.add_paragraph("中英混排测试标题")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(18)
    document.add_paragraph("1 测试范围")
    paragraph = document.add_paragraph()
    paragraph.add_run("正文使用中文字体，并包含 English words and numbers 2026。")
    emphasis = paragraph.add_run("这段强调必须保留。")
    emphasis.bold = True
    second = document.add_paragraph("第二段 intentionally uses mixed language for run-font checks.")
    second.paragraph_format.first_line_indent = Mm(3)
    _save(document, path)


def _lists(path: Path) -> None:
    document = _document()
    document.add_paragraph("1 列表布局")
    document.add_paragraph("项目符号中文 Chinese item", style="List Bullet")
    document.add_paragraph("第二个项目符号", style="List Bullet")
    document.add_paragraph("编号条目 one", style="List Number")
    document.add_paragraph("编号条目 two", style="List Number")
    document.add_paragraph("列表后的普通正文 paragraph after lists。")
    _save(document, path)


def _long_numbered_sentence(path: Path) -> None:
    document = _document()
    document.add_paragraph("1 研究背景")
    document.add_paragraph(
        "1.1 现状分析 当前 Word 文档存在大量不规范样式：有的正文用二号字，"
        "有的标题用小五字号，中英文混排字体不区分，段落有无缩进混杂，空行杂乱无章。"
    )
    document.add_paragraph("1.2 真正的二级标题")
    document.add_paragraph("标题后的普通正文用于确认角色边界。")
    _save(document, path)


def _sections_and_tables(path: Path) -> None:
    document = _document()
    first = document.sections[0]
    first.left_margin = Mm(24)
    first.right_margin = Mm(24)
    document.add_paragraph("1 表格与节保护")
    document.add_paragraph("正文 mixed content before the table。")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "字段"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "中文内容"
    table.cell(1, 1).text = "English content"
    image_paragraph = document.add_paragraph()
    image_paragraph.add_run().add_picture(BytesIO(PNG_1X1), width=Mm(8))
    second = document.add_section(WD_SECTION.NEW_PAGE)
    second.orientation = WD_ORIENT.LANDSCAPE
    second.page_width, second.page_height = second.page_height, second.page_width
    document.add_paragraph("2 横向节")
    document.add_paragraph("Landscape section body text 横向节正文。")
    _append_bookmark(document.paragraphs[-1], "scope_fixture_anchor", "7")
    _save(document, path)


def _append_bookmark(paragraph: Paragraph, name: str, bookmark_id: str) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), bookmark_id)
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), bookmark_id)
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--outdir",
        type=Path,
        default=Path("tests/fixtures/natural-language-scope"),
    )
    args = parser.parse_args()
    builders = {
        "mixed-body.docx": _mixed_body,
        "numbered-lists.docx": _lists,
        "long-numbered-sentence.docx": _long_numbered_sentence,
        "sections-and-tables.docx": _sections_and_tables,
    }
    for filename, builder in builders.items():
        builder(args.outdir / filename)


if __name__ == "__main__":
    main()
