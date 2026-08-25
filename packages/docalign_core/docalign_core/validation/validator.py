from __future__ import annotations

from pathlib import Path
from typing import Any
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn
from lxml import etree

from docalign_core.docx.formatting import (
    effective_table_font_size,
    iter_all_runs,
    run_is_protected,
)
from docalign_core.docx.parser import iter_block_items, parse_docx
from docalign_core.docx.safety import DocxSafetyError
from docalign_core.domain.audit import ValidationIssue, ValidationReport
from docalign_core.domain.document_ir import DocumentIR, ParagraphIR, TableIR
from docalign_core.domain.enums import SemanticRole, Severity
from docalign_core.domain.formatting_spec import (
    FontSpec,
    FormattingProperty,
    FormattingSpec,
    resolve_role_spec,
)
from docalign_core.engine.planner import style_name_for


class DocumentValidator:
    def validate(
        self,
        output_path: Path,
        spec: FormattingSpec,
        expected_document: DocumentIR,
    ) -> ValidationReport:
        issues: list[ValidationIssue] = []
        try:
            actual = parse_docx(output_path, document_id=expected_document.document_id)
        except DocxSafetyError as exc:
            return ValidationReport(
                valid=False,
                issues=[
                    ValidationIssue(
                        code=exc.code,
                        severity=Severity.FATAL,
                        message=exc.message,
                        details=exc.details,
                    )
                ],
                content_fingerprint_before=expected_document.content_fingerprint.digest,
            )
        except Exception as exc:  # mapped to a stable fatal boundary
            return ValidationReport(
                valid=False,
                issues=[
                    ValidationIssue(
                        code="OUTPUT_VALIDATION_FAILED",
                        severity=Severity.FATAL,
                        message=f"The output DOCX could not be reopened: {type(exc).__name__}.",
                    )
                ],
                content_fingerprint_before=expected_document.content_fingerprint.digest,
            )

        self._validate_content(expected_document, actual, spec, issues)
        self._validate_structure(expected_document, actual, issues)
        self._validate_formatting(output_path, expected_document, actual, spec, issues)
        self._validate_visual_cleanup(output_path, spec, issues)
        valid = not any(issue.severity in {Severity.FATAL, Severity.ERROR} for issue in issues)
        return ValidationReport(
            valid=valid,
            issues=issues,
            content_fingerprint_before=expected_document.content_fingerprint.digest,
            content_fingerprint_after=actual.content_fingerprint.digest,
        )

    def _validate_content(
        self,
        expected: DocumentIR,
        actual: DocumentIR,
        spec: FormattingSpec,
        issues: list[ValidationIssue],
    ) -> None:
        before = expected.content_fingerprint
        after = actual.content_fingerprint
        comparisons = (
            ("paragraph_texts", before.paragraph_texts, after.paragraph_texts),
            ("table_cell_texts", before.table_cell_texts, after.table_cell_texts),
            ("header_footer_texts", before.header_footer_texts, after.header_footer_texts),
            ("bookmark_names", before.bookmark_names, after.bookmark_names),
            ("image_hashes", before.image_hashes, after.image_hashes),
            (
                "relationship_signatures",
                before.relationship_signatures,
                after.relationship_signatures,
            ),
            (
                "unsupported_block_signatures",
                before.unsupported_block_signatures,
                after.unsupported_block_signatures,
            ),
        )
        for field, left, right in comparisons:
            if left != right:
                issues.append(
                    ValidationIssue(
                        code="CONTENT_INTEGRITY_FAILED",
                        severity=Severity.FATAL,
                        message=f"Content fingerprint component changed unexpectedly: {field}.",
                        details={"component": field},
                    )
                )
        if not _field_instructions_preserved(
            before.field_instructions,
            after.field_instructions,
            allow_page=bool(spec.page_numbers and spec.page_numbers.enabled),
        ):
            issues.append(
                ValidationIssue(
                    code="CONTENT_INTEGRITY_FAILED",
                    severity=Severity.FATAL,
                    message="Existing field instructions changed unexpectedly.",
                )
            )

    def _validate_structure(
        self,
        expected: DocumentIR,
        actual: DocumentIR,
        issues: list[ValidationIssue],
    ) -> None:
        pairs = (
            (
                "section",
                expected.content_fingerprint.section_count,
                actual.content_fingerprint.section_count,
            ),
            (
                "table",
                expected.content_fingerprint.table_count,
                actual.content_fingerprint.table_count,
            ),
            (
                "image",
                expected.content_fingerprint.image_count,
                actual.content_fingerprint.image_count,
            ),
        )
        for name, before, after in pairs:
            if before != after:
                issues.append(
                    ValidationIssue(
                        code="CONTENT_INTEGRITY_FAILED",
                        severity=Severity.FATAL,
                        message=f"The {name} count changed from {before} to {after}.",
                    )
                )
        if expected.content_fingerprint.block_kinds != actual.content_fingerprint.block_kinds:
            issues.append(
                ValidationIssue(
                    code="CONTENT_INTEGRITY_FAILED",
                    severity=Severity.FATAL,
                    message="The ordered top-level block structure changed unexpectedly.",
                )
            )
        expected_binary = _binary_part_hashes(expected)
        actual_binary = _binary_part_hashes(actual)
        if expected_binary != actual_binary:
            issues.append(
                ValidationIssue(
                    code="CONTENT_INTEGRITY_FAILED",
                    severity=Severity.FATAL,
                    message="One or more non-XML package parts changed unexpectedly.",
                    details={
                        "missing_or_changed": sorted(
                            path
                            for path, digest in expected_binary.items()
                            if actual_binary.get(path) != digest
                        ),
                        "unexpected": sorted(actual_binary.keys() - expected_binary.keys()),
                    },
                )
            )
        expected_protected_xml = _protected_xml_part_hashes(expected)
        actual_protected_xml = _protected_xml_part_hashes(actual)
        if expected_protected_xml != actual_protected_xml:
            issues.append(
                ValidationIssue(
                    code="PROTECTED_PACKAGE_PART_CHANGED",
                    severity=Severity.FATAL,
                    message=(
                        "A protected comments, footnotes, endnotes, or custom XML part changed."
                    ),
                    details={
                        "missing_or_changed": sorted(
                            path
                            for path, digest in expected_protected_xml.items()
                            if actual_protected_xml.get(path) != digest
                        ),
                        "unexpected": sorted(
                            actual_protected_xml.keys() - expected_protected_xml.keys()
                        ),
                    },
                )
            )

    def _validate_formatting(
        self,
        output_path: Path,
        expected: DocumentIR,
        actual: DocumentIR,
        spec: FormattingSpec,
        issues: list[ValidationIssue],
    ) -> None:
        actual_by_index = {block.index: block for block in actual.blocks}
        for block in expected.blocks:
            if not isinstance(block, ParagraphIR):
                continue
            if block.is_empty:
                continue
            role = block.detected_role
            if role == SemanticRole.UNKNOWN:
                if spec.baseline is None and not spec.behavior.apply_to_unknown_roles:
                    continue
                if spec.baseline is None:
                    role = spec.behavior.unknown_role_fallback
            role_spec = resolve_role_spec(spec, role)
            if role_spec is None:
                continue
            candidate = actual_by_index.get(block.index)
            if not isinstance(candidate, ParagraphIR):
                issues.append(
                    ValidationIssue(
                        code="DOCUMENT_STRUCTURE_MISMATCH",
                        severity=Severity.ERROR,
                        node_id=block.node_id,
                        locator=block.locator,
                        message=(
                            "A formatted paragraph no longer occupies the expected block position."
                        ),
                    )
                )
                continue
            expected_style = style_name_for(role, role_spec.style_name)
            if block.numbering is None and candidate.current_style_name != expected_style:
                issues.append(
                    ValidationIssue(
                        code="OUTPUT_VALIDATION_FAILED",
                        severity=Severity.ERROR,
                        node_id=block.node_id,
                        locator=block.locator,
                        message=(
                            f"Expected style {expected_style}, "
                            f"got {candidate.current_style_name!r}."
                        ),
                    )
                )
            if role_spec.font is None:
                continue
            for run in candidate.runs:
                if run.protected or not run.text:
                    continue
                font = run.formatting
                expected_fonts = {
                    "east_asia": role_spec.font.east_asia,
                    "ascii": role_spec.font.ascii,
                    "high_ansi": role_spec.font.high_ansi or role_spec.font.ascii,
                    "complex_script": (
                        role_spec.font.complex_script
                        or role_spec.font.high_ansi
                        or role_spec.font.ascii
                    ),
                }
                actual_fonts = {
                    "east_asia": font.east_asia_font,
                    "ascii": font.ascii_font,
                    "high_ansi": font.high_ansi_font,
                    "complex_script": font.complex_script_font,
                }
                mismatches: dict[str, dict[str, object]] = {
                    key: {"expected": value, "actual": actual_fonts[key]}
                    for key, value in expected_fonts.items()
                    if value is not None and actual_fonts[key] != value
                }
                if role_spec.font.size_pt is not None and font.size_pt != role_spec.font.size_pt:
                    mismatches["size_pt"] = {
                        "expected": role_spec.font.size_pt,
                        "actual": font.size_pt,
                    }
                for property_name, expected_value, actual_value in (
                    (
                        FormattingProperty.FONT_BOLD,
                        role_spec.font.bold,
                        font.bold,
                    ),
                    (
                        FormattingProperty.FONT_ITALIC,
                        role_spec.font.italic,
                        font.italic,
                    ),
                    (
                        FormattingProperty.FONT_UNDERLINE,
                        role_spec.font.underline,
                        font.underline,
                    ),
                    (
                        FormattingProperty.FONT_COLOR,
                        role_spec.font.color_hex,
                        font.color_hex,
                    ),
                ):
                    if (
                        property_name in role_spec.force.properties
                        and expected_value is not None
                        and actual_value != expected_value
                    ):
                        mismatches[property_name.value] = {
                            "expected": expected_value,
                            "actual": actual_value,
                        }
                if mismatches:
                    issues.append(
                        ValidationIssue(
                            code="OUTPUT_VALIDATION_FAILED",
                            severity=Severity.ERROR,
                            node_id=block.node_id,
                            locator=block.locator,
                            message="Run font properties do not match the role specification.",
                            details={"run_id": run.run_id, "mismatches": mismatches},
                        )
                    )
        self._validate_tables(output_path, expected, spec, issues)

    def _validate_tables(
        self,
        output_path: Path,
        expected: DocumentIR,
        spec: FormattingSpec,
        issues: list[ValidationIssue],
    ) -> None:
        table_spec = spec.tables
        if table_spec is None:
            return
        document = Document(str(output_path))
        actual_blocks = list(iter_block_items(document))
        for block in expected.blocks:
            if not isinstance(block, TableIR):
                continue
            if block.index >= len(actual_blocks) or actual_blocks[block.index][0] != "table":
                issues.append(
                    ValidationIssue(
                        code="DOCUMENT_STRUCTURE_MISMATCH",
                        severity=Severity.ERROR,
                        node_id=block.node_id,
                        locator=block.locator,
                        message="A formatted table no longer occupies the expected block position.",
                    )
                )
                continue
            table = actual_blocks[block.index][1]
            table_mismatches: dict[str, dict[str, object]] = {}
            if table_spec.alignment is not None:
                actual_alignment = _enum_name(table.alignment)
                if actual_alignment != table_spec.alignment.value:
                    table_mismatches["alignment"] = {
                        "expected": table_spec.alignment.value,
                        "actual": actual_alignment,
                    }
            if table_spec.grid_borders and table._tbl.tblPr.find(qn("w:tblBorders")) is None:
                table_mismatches["grid_borders"] = {"expected": True, "actual": False}
            if table_mismatches:
                issues.append(
                    ValidationIssue(
                        code="TABLE_FORMAT_VALIDATION_FAILED",
                        severity=Severity.ERROR,
                        node_id=block.node_id,
                        locator=block.locator,
                        message="Table properties do not match the formatting specification.",
                        details={"mismatches": table_mismatches},
                    )
                )

            expected_font = table_spec.font
            expected_size = effective_table_font_size(
                expected_font.size_pt if expected_font else None,
                len(table.columns),
                adaptive=table_spec.adaptive_font_size,
                minimum_size_pt=table_spec.min_font_size_pt,
            )
            seen_cells: set[object] = set()
            for row_index, row in enumerate(table.rows):
                row_properties = row._tr.trPr
                row_locator = f"{block.locator}.r{row_index + 1}"
                if row_index == 0 and table_spec.repeat_header_row:
                    header = (
                        row_properties.find(qn("w:tblHeader"))
                        if row_properties is not None
                        else None
                    )
                    if header is None or not _ooxml_boolean_is_on(header):
                        issues.append(
                            ValidationIssue(
                                code="TABLE_HEADER_REPEAT_MISSING",
                                severity=Severity.ERROR,
                                node_id=block.node_id,
                                locator=row_locator,
                                message=(
                                    "The first table row is not configured to repeat as a header."
                                ),
                            )
                        )
                if table_spec.prevent_row_split:
                    cant_split = (
                        row_properties.find(qn("w:cantSplit"))
                        if row_properties is not None
                        else None
                    )
                    if cant_split is None or not _ooxml_boolean_is_on(cant_split):
                        issues.append(
                            ValidationIssue(
                                code="TABLE_ROW_SPLIT_RULE_MISSING",
                                severity=Severity.ERROR,
                                node_id=block.node_id,
                                locator=row_locator,
                                message="A table row may split across pages contrary to the spec.",
                            )
                        )
                for column_index, cell in enumerate(row.cells):
                    if cell._tc in seen_cells:
                        continue
                    seen_cells.add(cell._tc)
                    cell_locator = f"{row_locator}.c{column_index + 1}"
                    if table_spec.cell_vertical_alignment is not None:
                        actual_vertical = _enum_name(cell.vertical_alignment)
                        if actual_vertical != table_spec.cell_vertical_alignment:
                            issues.append(
                                ValidationIssue(
                                    code="TABLE_CELL_ALIGNMENT_FAILED",
                                    severity=Severity.ERROR,
                                    node_id=block.node_id,
                                    locator=cell_locator,
                                    message="Table cell vertical alignment is not compliant.",
                                    details={
                                        "expected": table_spec.cell_vertical_alignment,
                                        "actual": actual_vertical,
                                    },
                                )
                            )
                    if expected_font is None:
                        continue
                    for paragraph_index, paragraph in enumerate(cell.paragraphs):
                        for run_index, run in enumerate(iter_all_runs(paragraph)):
                            if run_is_protected(run) or not run.text:
                                continue
                            mismatches = _table_run_font_mismatches(
                                run,
                                expected_font,
                                expected_size,
                                preserve_inline_emphasis=spec.behavior.preserve_inline_emphasis,
                                cleanup_color=(
                                    spec.visual_cleanup.text_color_hex
                                    if spec.visual_cleanup is not None
                                    else None
                                ),
                            )
                            if not mismatches:
                                continue
                            issues.append(
                                ValidationIssue(
                                    code="TABLE_FONT_VALIDATION_FAILED",
                                    severity=Severity.ERROR,
                                    node_id=block.node_id,
                                    locator=(
                                        f"{cell_locator}.p{paragraph_index + 1}.r{run_index + 1}"
                                    ),
                                    message=(
                                        "Table-cell font properties do not match the table spec."
                                    ),
                                    details={"mismatches": mismatches},
                                )
                            )

    def _validate_visual_cleanup(
        self,
        output_path: Path,
        spec: FormattingSpec,
        issues: list[ValidationIssue],
    ) -> None:
        cleanup = spec.visual_cleanup
        if cleanup is None:
            return
        counts = _visual_cleanup_violations(output_path, spec)
        if any(counts.values()):
            issues.append(
                ValidationIssue(
                    code="VISUAL_CLEANUP_FAILED",
                    severity=Severity.ERROR,
                    message="One or more requested document-wide visual cleanup rules remain.",
                    details={"violations": counts},
                )
            )


def _field_instructions_preserved(
    before: list[str],
    after: list[str],
    *,
    allow_page: bool,
) -> bool:
    remaining = list(after)
    for value in before:
        try:
            remaining.remove(value)
        except ValueError:
            return False
    if not remaining:
        return True
    return allow_page and all(value.strip().upper() == "PAGE" for value in remaining)


def _binary_part_hashes(document: DocumentIR) -> dict[str, str]:
    return {
        part.path: part.sha256
        for part in document.package_parts
        if not part.path.lower().endswith((".xml", ".rels"))
    }


def _protected_xml_part_hashes(document: DocumentIR) -> dict[str, str]:
    protected_parts = {
        "word/footnotes.xml",
        "word/endnotes.xml",
        "word/people.xml",
    }
    return {
        part.path: part.sha256
        for part in document.package_parts
        if (
            part.path in protected_parts
            or part.path.startswith("customXml/")
            or part.path.startswith("word/comments")
            or part.path.startswith("word/glossary/")
        )
        and part.path.lower().endswith((".xml", ".rels"))
    }


def _table_run_font_mismatches(
    run: Any,
    expected_font: FontSpec,
    expected_size: float | None,
    *,
    preserve_inline_emphasis: bool,
    cleanup_color: str | None,
) -> dict[str, dict[str, object]]:
    run_properties = run._r.rPr
    fonts = run_properties.find(qn("w:rFonts")) if run_properties is not None else None
    actual_fonts = {
        "east_asia": fonts.get(qn("w:eastAsia")) if fonts is not None else None,
        "ascii": fonts.get(qn("w:ascii")) if fonts is not None else None,
        "high_ansi": fonts.get(qn("w:hAnsi")) if fonts is not None else None,
        "complex_script": fonts.get(qn("w:cs")) if fonts is not None else None,
    }
    expected_fonts = {
        "east_asia": expected_font.east_asia,
        "ascii": expected_font.ascii,
        "high_ansi": expected_font.high_ansi or expected_font.ascii,
        "complex_script": (
            expected_font.complex_script or expected_font.high_ansi or expected_font.ascii
        ),
    }
    mismatches = {
        key: {"expected": value, "actual": actual_fonts[key]}
        for key, value in expected_fonts.items()
        if value is not None and actual_fonts[key] != value
    }
    actual_size = round(run.font.size.pt, 3) if run.font.size is not None else None
    if expected_size is not None and actual_size != expected_size:
        mismatches["size_pt"] = {"expected": expected_size, "actual": actual_size}
    if not preserve_inline_emphasis:
        for key, expected, actual in (
            ("bold", expected_font.bold, run.bold),
            ("italic", expected_font.italic, run.italic),
            (
                "underline",
                expected_font.underline,
                bool(run.underline) if run.underline is not None else None,
            ),
        ):
            if expected is not None and actual != expected:
                mismatches[key] = {"expected": expected, "actual": actual}
    if cleanup_color is not None:
        color = run.font.color.rgb
        actual_color = str(color) if color is not None else None
        if actual_color != cleanup_color:
            mismatches["color_hex"] = {
                "expected": cleanup_color,
                "actual": actual_color,
            }
    return mismatches


def _ooxml_boolean_is_on(element: Any) -> bool:
    return (element.get(qn("w:val")) or "1").lower() not in {"0", "false", "off"}


def _enum_name(value: Any) -> str | None:
    if value is None:
        return None
    name = getattr(value, "name", None)
    return str(name).lower() if name else str(value).lower()


def _visual_cleanup_violations(output_path: Path, spec: FormattingSpec) -> dict[str, int]:
    cleanup = spec.visual_cleanup
    if cleanup is None:
        return {}
    counts = {
        "text_color": 0,
        "text_highlight": 0,
        "character_shading": 0,
        "paragraph_shading": 0,
        "table_cell_shading": 0,
        "page_background": 0,
    }
    with ZipFile(output_path) as package:
        for name in package.namelist():
            if not name.startswith("word/") or not name.endswith(".xml"):
                continue
            try:
                root = etree.fromstring(package.read(name))
            except etree.XMLSyntaxError:
                continue
            if cleanup.text_color_hex is not None:
                for color in root.iter(qn("w:color")):
                    if color.get(qn("w:val")) != cleanup.text_color_hex or any(
                        attribute in color.attrib
                        for attribute in (
                            qn("w:themeColor"),
                            qn("w:themeTint"),
                            qn("w:themeShade"),
                        )
                    ):
                        counts["text_color"] += 1
                for run in root.iter(qn("w:r")):
                    if not any((text.text or "") for text in run.iter(qn("w:t"))):
                        continue
                    run_properties = run.find(qn("w:rPr"))
                    run_color = (
                        run_properties.find(qn("w:color")) if run_properties is not None else None
                    )
                    if run_color is None or run_color.get(qn("w:val")) != cleanup.text_color_hex:
                        counts["text_color"] += 1
            if cleanup.remove_text_highlight:
                counts["text_highlight"] += sum(
                    1
                    for properties in root.iter(qn("w:rPr"))
                    if properties.find(qn("w:highlight")) is not None
                )
            for enabled, key, parent_name in (
                (cleanup.remove_character_shading, "character_shading", "w:rPr"),
                (cleanup.remove_paragraph_shading, "paragraph_shading", "w:pPr"),
                (cleanup.remove_table_cell_shading, "table_cell_shading", "w:tcPr"),
            ):
                if enabled:
                    counts[key] += sum(
                        1
                        for properties in root.iter(qn(parent_name))
                        if properties.find(qn("w:shd")) is not None
                    )
            if cleanup.remove_page_background:
                counts["page_background"] += sum(
                    1
                    for document in root.iter(qn("w:document"))
                    if document.find(qn("w:background")) is not None
                )
    return counts
