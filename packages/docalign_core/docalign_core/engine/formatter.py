from __future__ import annotations

import os
import re
from pathlib import Path
from typing import Any
from zipfile import ZipFile

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Length, Mm
from lxml import etree
from pydantic import Field

from docalign_core.docx.formatting import (
    ALIGNMENT_MAP,
    CELL_ALIGNMENT_MAP,
    TABLE_ALIGNMENT_MAP,
    apply_direct_paragraph_spec,
    clear_covered_paragraph_format,
    effective_table_font_size,
    ensure_portable_font_aliases,
    ensure_role_style,
    insert_page_field,
    iter_all_runs,
    mm_to_twips,
    run_is_protected,
    set_repeat_table_header,
    set_row_cant_split,
    set_run_font,
    set_table_grid_borders,
    set_table_width,
    suggest_table_column_weights,
)
from docalign_core.docx.parser import iter_block_items
from docalign_core.docx.safety import sha256_file
from docalign_core.domain.audit import FormattingPlan, MutationRecord, PlanWarning
from docalign_core.domain.base import StrictModel
from docalign_core.domain.document_ir import DocumentIR, ParagraphIR, TableIR
from docalign_core.domain.enums import SemanticRole
from docalign_core.domain.formatting_spec import (
    FontSpec,
    FormattingProperty,
    FormattingSpec,
    Orientation,
    PageSize,
    ParagraphSpec,
    TableWidthPolicy,
    VisualCleanupSpec,
    resolve_role_spec,
)
from docalign_core.engine.planner import configured_role_specs, style_name_for


class FormattingExecutionResult(StrictModel):
    output_path: str
    mutations: list[MutationRecord] = Field(default_factory=list)
    warnings: list[PlanWarning] = Field(default_factory=list)


class FormattingEngine:
    def apply(
        self,
        source_path: Path,
        document_ir: DocumentIR,
        spec: FormattingSpec,
        plan: FormattingPlan,
        output_path: Path,
    ) -> FormattingExecutionResult:
        if sha256_file(source_path) != document_ir.source_sha256:
            raise ValueError("ANALYSIS_SOURCE_MISMATCH")
        document = Document(str(source_path))
        mutations: list[MutationRecord] = []
        warnings = list(plan.warnings)
        operation_lookup = {
            (operation.node_id, operation.operation_type.value): operation.operation_id
            for operation in plan.operations
        }

        self._apply_sections(document, spec, mutations, operation_lookup)
        styles: dict[SemanticRole, Any] = {}
        for role, role_spec in configured_role_specs(document_ir, spec).items():
            name = style_name_for(role, role_spec.style_name)
            before = name in {style.name for style in document.styles}
            styles[role] = ensure_role_style(document, name, role_spec)
            mutations.append(
                MutationRecord(
                    operation_id=operation_lookup.get((None, "create_or_update_style"), "style"),
                    property_path=f"styles.{name}",
                    before="existing" if before else None,
                    after="configured",
                    status="changed" if not before else "already_compliant",
                )
            )

        actual_blocks = list(iter_block_items(document))
        for block in document_ir.blocks:
            if block.index >= len(actual_blocks):
                warnings.append(
                    PlanWarning(
                        code="DOCUMENT_STRUCTURE_MISMATCH",
                        node_id=block.node_id,
                        locator=block.locator,
                        message="The source block could not be located during formatting.",
                    )
                )
                continue
            kind, actual = actual_blocks[block.index]
            if isinstance(block, ParagraphIR) and kind == "paragraph":
                self._format_paragraph(
                    actual,
                    block,
                    spec,
                    styles,
                    mutations,
                    warnings,
                    operation_lookup,
                )
            elif isinstance(block, TableIR) and kind == "table" and spec.tables is not None:
                self._format_table(
                    actual,
                    block,
                    document,
                    spec,
                    mutations,
                    warnings,
                    operation_lookup,
                )

        self._format_headers_footers(document, spec, mutations, warnings, operation_lookup)
        ensure_portable_font_aliases(document, _configured_font_names(spec))
        output_path.parent.mkdir(parents=True, exist_ok=True)
        document.save(str(output_path))
        self._normalize_saved_package_visuals(output_path, spec, mutations, operation_lookup)
        return FormattingExecutionResult(
            output_path=str(output_path), mutations=mutations, warnings=_dedupe(warnings)
        )

    def _apply_sections(
        self,
        document: Any,
        spec: FormattingSpec,
        mutations: list[MutationRecord],
        operation_lookup: dict[tuple[str | None, str], str],
    ) -> None:
        if spec.document is None:
            return
        page = spec.document.page
        for index, section in enumerate(document.sections):
            before = {
                "width": section.page_width.twips if section.page_width else None,
                "height": section.page_height.twips if section.page_height else None,
                "orientation": getattr(section.orientation, "name", str(section.orientation)),
                "margins": [
                    value.twips if value else None
                    for value in (
                        section.top_margin,
                        section.bottom_margin,
                        section.left_margin,
                        section.right_margin,
                    )
                ],
            }
            currently_landscape = section.orientation == WD_ORIENT.LANDSCAPE
            preserve_landscape = (
                currently_landscape
                and page.preserve_existing_landscape_sections
                and not page.force_orientation_all_sections
                and page.orientation == Orientation.PORTRAIT
            )
            target_orientation = (
                Orientation.LANDSCAPE
                if preserve_landscape
                else page.orientation
                or (Orientation.LANDSCAPE if currently_landscape else Orientation.PORTRAIT)
            )
            if page.size == PageSize.A4:
                short_side: Length = Mm(210)
                long_side: Length = Mm(297)
            elif page.size == PageSize.LETTER:
                short_side = Inches(8.5)
                long_side = Inches(11)
            else:
                current_width = section.page_width or Mm(210)
                current_height = section.page_height or Mm(297)
                short_side = Length(min(current_width, current_height))
                long_side = Length(max(current_width, current_height))
            if page.size is not None or page.orientation is not None:
                if target_orientation == Orientation.LANDSCAPE:
                    section.orientation = WD_ORIENT.LANDSCAPE
                    section.page_width, section.page_height = long_side, short_side
                else:
                    section.orientation = WD_ORIENT.PORTRAIT
                    section.page_width, section.page_height = short_side, long_side
            if page.margin_top_mm is not None:
                section.top_margin = Mm(page.margin_top_mm)
            if page.margin_bottom_mm is not None:
                section.bottom_margin = Mm(page.margin_bottom_mm)
            if page.margin_left_mm is not None:
                section.left_margin = Mm(page.margin_left_mm)
            if page.margin_right_mm is not None:
                section.right_margin = Mm(page.margin_right_mm)
            if page.header_distance_mm is not None:
                section.header_distance = Mm(page.header_distance_mm)
            if page.footer_distance_mm is not None:
                section.footer_distance = Mm(page.footer_distance_mm)
            if spec.page_numbers and spec.page_numbers.start_at is not None:
                _set_page_number_start(section, spec.page_numbers.start_at)
            after = {
                "width": section.page_width.twips if section.page_width else None,
                "height": section.page_height.twips if section.page_height else None,
                "orientation": getattr(section.orientation, "name", str(section.orientation)),
                "margins": [
                    value.twips if value else None
                    for value in (
                        section.top_margin,
                        section.bottom_margin,
                        section.left_margin,
                        section.right_margin,
                    )
                ],
            }
            mutations.append(
                MutationRecord(
                    operation_id=operation_lookup.get(
                        (f"section-{index}", "set_section_layout"), f"section-{index}"
                    ),
                    node_id=f"section-{index}",
                    locator=f"s{index + 1}",
                    property_path="section.layout",
                    before=before,
                    after=after,
                    status="changed" if before != after else "already_compliant",
                )
            )

    def _format_paragraph(
        self,
        paragraph: Any,
        block: ParagraphIR,
        spec: FormattingSpec,
        styles: dict[SemanticRole, Any],
        mutations: list[MutationRecord],
        warnings: list[PlanWarning],
        operation_lookup: dict[tuple[str | None, str], str],
    ) -> None:
        if block.is_empty:
            return
        initial_alignment = str(paragraph.alignment)
        role = block.detected_role
        if role == SemanticRole.UNKNOWN:
            if spec.baseline is None and not spec.behavior.apply_to_unknown_roles:
                return
            if spec.baseline is None:
                role = spec.behavior.unknown_role_fallback
        role_spec = resolve_role_spec(spec, role)
        if role_spec is None:
            return
        before_style = paragraph.style.name if paragraph.style is not None else None
        if block.numbering is not None:
            list_paragraph_spec = _list_safe_paragraph_spec(role_spec.paragraph)
            if list_paragraph_spec is not None:
                if spec.behavior.normalize_direct_paragraph_formatting:
                    clear_covered_paragraph_format(paragraph, list_paragraph_spec)
                apply_direct_paragraph_spec(paragraph, list_paragraph_spec)
            warnings.append(
                PlanWarning(
                    code="NUMBERING_LAYOUT_PRESERVED",
                    node_id=block.node_id,
                    locator=block.locator,
                    message=(
                        "Existing list numbering, paragraph style, alignment, and indents were "
                        "preserved; compatible font and spacing rules were applied directly."
                    ),
                )
            )
        else:
            if (
                role_spec.paragraph is not None
                and spec.behavior.normalize_direct_paragraph_formatting
            ):
                clear_covered_paragraph_format(paragraph, role_spec.paragraph)
            paragraph.style = styles[role]
            if spec.behavior.preserve_right_aligned_signatures and _is_signature_block(block):
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        after_style = paragraph.style.name if paragraph.style is not None else None
        mutations.append(
            MutationRecord(
                operation_id=operation_lookup.get(
                    (block.node_id, "assign_paragraph_style"), f"style-{block.node_id}"
                ),
                node_id=block.node_id,
                locator=block.locator,
                property_path="paragraph.style",
                before=before_style,
                after=after_style,
                status=(
                    "skipped"
                    if block.numbering is not None
                    else "changed"
                    if before_style != after_style
                    else "already_compliant"
                ),
            )
        )
        if role_spec.font is not None and spec.behavior.normalize_direct_run_formatting:
            for run_index, run in enumerate(iter_all_runs(paragraph)):
                if run_is_protected(run):
                    warnings.append(
                        PlanWarning(
                            code="PROTECTED_RUN_SKIPPED",
                            node_id=block.node_id,
                            locator=block.locator,
                            message=(
                                "A protected run was preserved without direct font normalization."
                            ),
                        )
                    )
                    continue
                if not run.text:
                    continue
                before = _font_state(run)
                set_run_font(
                    run,
                    role_spec.font,
                    _effective_font_force(
                        spec,
                        role_spec.font,
                        role_spec.force.properties,
                    ),
                )
                after = _font_state(run)
                mutations.append(
                    MutationRecord(
                        operation_id=operation_lookup.get(
                            (block.node_id, "set_run_font"), f"font-{block.node_id}"
                        ),
                        node_id=block.node_id,
                        locator=f"{block.locator}.r{run_index + 1}",
                        property_path=f"runs.{run_index}.font",
                        before=before,
                        after=after,
                        status="changed" if before != after else "already_compliant",
                    )
                )
        if (
            block.contains_drawing
            and not block.text.strip()
            and spec.figures
            and spec.figures.center_image_only_paragraphs
        ):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            after_alignment = str(paragraph.alignment)
            mutations.append(
                MutationRecord(
                    operation_id=operation_lookup.get(
                        (block.node_id, "align_image_paragraph"), f"image-{block.node_id}"
                    ),
                    node_id=block.node_id,
                    locator=block.locator,
                    property_path="paragraph.alignment",
                    before=initial_alignment,
                    after=after_alignment,
                    status=(
                        "changed" if initial_alignment != after_alignment else "already_compliant"
                    ),
                )
            )

    def _format_table(
        self,
        table: Any,
        block: TableIR,
        document: Any,
        spec: FormattingSpec,
        mutations: list[MutationRecord],
        warnings: list[PlanWarning],
        operation_lookup: dict[tuple[str | None, str], str],
    ) -> None:
        table_spec = spec.tables
        if table_spec is None:
            return
        before = {
            "alignment": str(table.alignment),
            "autofit": table.autofit,
        }
        if table_spec.alignment is not None and table_spec.alignment in TABLE_ALIGNMENT_MAP:
            table.alignment = TABLE_ALIGNMENT_MAP[table_spec.alignment]
        weights = suggest_table_column_weights(table) if table_spec.adaptive_column_widths else None
        minimum_width = mm_to_twips(table_spec.min_column_width_mm)
        if table_spec.width_policy == TableWidthPolicy.FIXED and table_spec.fixed_width_mm:
            set_table_width(
                table,
                mm_to_twips(table_spec.fixed_width_mm),
                column_weights=weights,
                min_column_width_twips=minimum_width,
            )
        elif table_spec.width_policy == TableWidthPolicy.FIT_PRINTABLE_WIDTH:
            section = _section_for_element(document, table._tbl)
            printable_width = int(
                section.page_width.twips - section.left_margin.twips - section.right_margin.twips
            )
            set_table_width(
                table,
                printable_width,
                column_weights=weights,
                min_column_width_twips=minimum_width,
            )
        if table_spec.grid_borders:
            set_table_grid_borders(table)
        effective_size = effective_table_font_size(
            table_spec.font.size_pt if table_spec.font else None,
            len(table.columns),
            adaptive=table_spec.adaptive_font_size,
            minimum_size_pt=table_spec.min_font_size_pt,
        )
        effective_font = (
            table_spec.font.model_copy(update={"size_pt": effective_size})
            if table_spec.font is not None
            else None
        )
        # Keep the XML elements themselves alive while walking the table.  Using
        # ``id(cell._tc)`` is unsafe with lxml: proxy objects can be released and
        # their Python ids reused between rows, which caused later cells to be
        # mistaken for merged-cell duplicates and left their fonts untouched.
        seen_cells: set[Any] = set()
        for row_index, row in enumerate(table.rows):
            if row_index == 0 and table_spec.repeat_header_row:
                set_repeat_table_header(row)
            if table_spec.prevent_row_split:
                set_row_cant_split(row)
            for cell in row.cells:
                cell_key = cell._tc
                if cell_key in seen_cells:
                    continue
                seen_cells.add(cell_key)
                if table_spec.cell_vertical_alignment is not None:
                    cell.vertical_alignment = CELL_ALIGNMENT_MAP[table_spec.cell_vertical_alignment]
                for paragraph in cell.paragraphs:
                    if table_spec.paragraph is not None:
                        apply_direct_paragraph_spec(paragraph, table_spec.paragraph)
                    if effective_font is not None:
                        for run in iter_all_runs(paragraph):
                            if run_is_protected(run):
                                warnings.append(
                                    PlanWarning(
                                        code="PROTECTED_RUN_SKIPPED",
                                        node_id=block.node_id,
                                        locator=block.locator,
                                        message="A protected table-cell run was preserved.",
                                    )
                                )
                            elif run.text:
                                set_run_font(
                                    run,
                                    effective_font,
                                    _effective_font_force(spec, effective_font, set()),
                                )
        after = {"alignment": str(table.alignment), "autofit": table.autofit}
        mutations.append(
            MutationRecord(
                operation_id=operation_lookup.get(
                    (block.node_id, "set_table_format"), f"table-{block.node_id}"
                ),
                node_id=block.node_id,
                locator=block.locator,
                property_path="table.format",
                before=before,
                after=after,
                status="changed" if before != after else "already_compliant",
            )
        )

    def _format_headers_footers(
        self,
        document: Any,
        spec: FormattingSpec,
        mutations: list[MutationRecord],
        warnings: list[PlanWarning],
        operation_lookup: dict[tuple[str | None, str], str],
    ) -> None:
        inserted_parts: set[str] = set()
        for section_index, section in enumerate(document.sections):
            if spec.page_numbers and spec.page_numbers.enabled:
                footer = section.footer
                part_key = str(footer.part.partname)
                if part_key not in inserted_parts:
                    inserted_parts.add(part_key)
                    has_page = any(
                        "PAGE" in (node.text or "").upper()
                        for node in footer._element.iter(qn("w:instrText"))
                    )
                    if not has_page:
                        target = footer.paragraphs[0]
                        if target.text.strip() or len(target._p) > 1:
                            target = footer.add_paragraph()
                        target.alignment = ALIGNMENT_MAP[spec.page_numbers.alignment]
                        changed = insert_page_field(target)
                        mutations.append(
                            MutationRecord(
                                operation_id=operation_lookup.get(
                                    (None, "insert_page_number"), "page-number"
                                ),
                                node_id=f"footer-{section_index}",
                                property_path="footer.page_number",
                                before=None,
                                after="PAGE",
                                status="changed" if changed else "already_compliant",
                            )
                        )
            for part_name, format_spec, containers in (
                (
                    "header",
                    spec.headers,
                    (section.header, section.first_page_header, section.even_page_header),
                ),
                (
                    "footer",
                    spec.footers,
                    (section.footer, section.first_page_footer, section.even_page_footer),
                ),
            ):
                if format_spec is None:
                    continue
                for container in containers:
                    for paragraph in container.paragraphs:
                        if format_spec.paragraph is not None:
                            apply_direct_paragraph_spec(paragraph, format_spec.paragraph)
                        if format_spec.font is not None:
                            for run in iter_all_runs(paragraph):
                                if not run_is_protected(run) and (
                                    run.text or _run_has_page_field(run)
                                ):
                                    set_run_font(
                                        run,
                                        format_spec.font,
                                        _effective_font_force(
                                            spec,
                                            format_spec.font,
                                            set(),
                                        ),
                                    )
                mutations.append(
                    MutationRecord(
                        operation_id=operation_lookup.get(
                            (None, f"format_{part_name}"), f"{part_name}-{section_index}"
                        ),
                        node_id=f"{part_name}-{section_index}",
                        property_path=f"{part_name}.format",
                        before=None,
                        after="configured",
                        status="changed",
                    )
                )

    def _normalize_saved_package_visuals(
        self,
        output_path: Path,
        spec: FormattingSpec,
        mutations: list[MutationRecord],
        operation_lookup: dict[tuple[str | None, str], str],
    ) -> None:
        cleanup = spec.visual_cleanup
        if cleanup is None:
            return
        operation_id = operation_lookup.get((None, "normalize_document_visuals"), "visual-cleanup")
        changed_by_property = _normalize_docx_package_visuals(output_path, cleanup)
        targets: dict[str, object] = {
            "text_color_hex": cleanup.text_color_hex,
            "remove_text_highlight": True,
            "remove_character_shading": True,
            "remove_paragraph_shading": True,
            "remove_table_cell_shading": True,
            "remove_page_background": True,
        }
        actions = [
            (property_name, changed, targets[property_name])
            for property_name, changed in changed_by_property.items()
        ]
        for property_name, changed, target in actions:
            mutations.append(
                MutationRecord(
                    operation_id=operation_id,
                    property_path=f"visual_cleanup.{property_name}",
                    before={"noncompliant_nodes": changed},
                    after=target,
                    status="changed" if changed else "already_compliant",
                )
            )


def atomic_promote(temp_path: Path, final_path: Path) -> None:
    final_path.parent.mkdir(parents=True, exist_ok=True)
    os.replace(temp_path, final_path)


def _configured_font_names(spec: FormattingSpec) -> set[str]:
    fonts: list[FontSpec | None] = [
        spec.baseline.font if spec.baseline else None,
        spec.tables.font if spec.tables else None,
        spec.headers.font if spec.headers else None,
        spec.footers.font if spec.footers else None,
    ]
    fonts.extend(role.font for role in spec.roles.values())
    names: set[str] = set()
    for font in fonts:
        if font is None:
            continue
        names.update(
            name
            for name in (
                font.east_asia,
                font.ascii,
                font.high_ansi,
                font.complex_script,
            )
            if name
        )
    return names


def _section_for_element(document: Any, target: Any) -> Any:
    """Return the section containing a top-level body element."""

    section_index = 0
    for element in document._element.body.iterchildren():
        if element is target:
            break
        if element.tag == qn("w:p") and element.find("./w:pPr/w:sectPr", element.nsmap) is not None:
            section_index += 1
    return document.sections[min(section_index, len(document.sections) - 1)]


def _font_state(run: Any) -> dict[str, object | None]:
    rpr = run._r.rPr
    fonts = rpr.find(qn("w:rFonts")) if rpr is not None else None
    return {
        "ascii": fonts.get(qn("w:ascii")) if fonts is not None else None,
        "hAnsi": fonts.get(qn("w:hAnsi")) if fonts is not None else None,
        "eastAsia": fonts.get(qn("w:eastAsia")) if fonts is not None else None,
        "cs": fonts.get(qn("w:cs")) if fonts is not None else None,
        "size_pt": run.font.size.pt if run.font.size is not None else None,
        "bold": run.bold,
        "italic": run.italic,
        "underline": bool(run.underline) if run.underline is not None else None,
    }


def _effective_font_force(
    spec: FormattingSpec,
    font: FontSpec,
    configured: set[FormattingProperty],
) -> set[FormattingProperty]:
    forced = set(configured)
    if spec.behavior.preserve_inline_emphasis:
        return forced
    for value, property_name in (
        (font.bold, FormattingProperty.FONT_BOLD),
        (font.italic, FormattingProperty.FONT_ITALIC),
        (font.underline, FormattingProperty.FONT_UNDERLINE),
    ):
        if value is not None:
            forced.add(property_name)
    return forced


def _set_page_number_start(section: Any, start_at: int) -> None:
    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num_type)
    pg_num_type.set(qn("w:start"), str(start_at))


def _run_has_page_field(run: Any) -> bool:
    return any("PAGE" in (node.text or "").upper() for node in run._r.iter(qn("w:instrText")))


def _dedupe(warnings: list[PlanWarning]) -> list[PlanWarning]:
    seen: set[tuple[str, str | None]] = set()
    result: list[PlanWarning] = []
    for warning in warnings:
        key = (warning.code, warning.node_id)
        if key not in seen:
            seen.add(key)
            result.append(warning)
    return result


def _normalize_visible_text_color(roots: list[Any], color_hex: str) -> int:
    changed = 0
    color_tag = qn("w:color")
    value_attr = qn("w:val")
    theme_attrs = (qn("w:themeColor"), qn("w:themeTint"), qn("w:themeShade"))
    for root in roots:
        for color in root.iter(color_tag):
            if color.get(value_attr) != color_hex or any(
                attribute in color.attrib for attribute in theme_attrs
            ):
                changed += 1
            color.set(value_attr, color_hex)
            for attribute in theme_attrs:
                color.attrib.pop(attribute, None)
        for run in root.iter(qn("w:r")):
            if not any((text.text or "") for text in run.iter(qn("w:t"))):
                continue
            run_properties = run.find(qn("w:rPr"))
            if run_properties is None:
                run_properties = OxmlElement("w:rPr")
                run.insert(0, run_properties)
            color = run_properties.find(color_tag)
            if color is None:
                color = OxmlElement("w:color")
                run_properties.append(color)
                changed += 1
            elif color.get(value_attr) != color_hex or any(
                attribute in color.attrib for attribute in theme_attrs
            ):
                changed += 1
            color.set(value_attr, color_hex)
            for attribute in theme_attrs:
                color.attrib.pop(attribute, None)
    return changed


def _remove_child_property(roots: list[Any], parent_name: str, child_name: str) -> int:
    removed = 0
    child_tag = qn(child_name)
    for root in roots:
        for parent in root.iter(qn(parent_name)):
            for child in list(parent.findall(child_tag)):
                parent.remove(child)
                removed += 1
    return removed


def _normalize_docx_package_visuals(
    output_path: Path, cleanup: VisualCleanupSpec
) -> dict[str, int]:
    enabled: dict[str, bool] = {
        "text_color_hex": cleanup.text_color_hex is not None,
        "remove_text_highlight": cleanup.remove_text_highlight,
        "remove_character_shading": cleanup.remove_character_shading,
        "remove_paragraph_shading": cleanup.remove_paragraph_shading,
        "remove_table_cell_shading": cleanup.remove_table_cell_shading,
        "remove_page_background": cleanup.remove_page_background,
    }
    changed = {name: 0 for name, is_enabled in enabled.items() if is_enabled}
    package_temp = output_path.with_name(f".{output_path.name}.visual-cleanup.tmp")
    try:
        with ZipFile(output_path, "r") as source, ZipFile(package_temp, "w") as destination:
            destination.comment = source.comment
            for info in source.infolist():
                data = source.read(info.filename)
                if info.filename.startswith("word/") and info.filename.endswith(".xml"):
                    try:
                        root = etree.fromstring(data)
                    except etree.XMLSyntaxError:
                        root = None
                    if root is not None:
                        part_changes = 0
                        if cleanup.text_color_hex is not None:
                            count = _normalize_visible_text_color([root], cleanup.text_color_hex)
                            changed["text_color_hex"] += count
                            part_changes += count
                        for is_enabled, property_name, parent_name, child_name in (
                            (
                                cleanup.remove_text_highlight,
                                "remove_text_highlight",
                                "w:rPr",
                                "w:highlight",
                            ),
                            (
                                cleanup.remove_character_shading,
                                "remove_character_shading",
                                "w:rPr",
                                "w:shd",
                            ),
                            (
                                cleanup.remove_paragraph_shading,
                                "remove_paragraph_shading",
                                "w:pPr",
                                "w:shd",
                            ),
                            (
                                cleanup.remove_table_cell_shading,
                                "remove_table_cell_shading",
                                "w:tcPr",
                                "w:shd",
                            ),
                            (
                                cleanup.remove_page_background,
                                "remove_page_background",
                                "w:document",
                                "w:background",
                            ),
                        ):
                            if is_enabled:
                                count = _remove_child_property([root], parent_name, child_name)
                                changed[property_name] += count
                                part_changes += count
                        if part_changes:
                            data = etree.tostring(
                                root,
                                encoding="UTF-8",
                                xml_declaration=True,
                                standalone=True,
                            )
                destination.writestr(info, data)
        os.replace(package_temp, output_path)
    finally:
        package_temp.unlink(missing_ok=True)
    return changed


def _list_safe_paragraph_spec(spec: ParagraphSpec | None) -> ParagraphSpec | None:
    if spec is None:
        return None
    payload = spec.model_dump(mode="json", exclude_none=True)
    for property_name in (
        "alignment",
        "first_line_indent_pt",
        "hanging_indent_pt",
        "left_indent_pt",
        "right_indent_pt",
        "keep_with_next",
        "page_break_before",
    ):
        payload.pop(property_name, None)
    return ParagraphSpec.model_validate(payload) if payload else None


def _is_signature_block(block: ParagraphIR) -> bool:
    alignment = (block.formatting.alignment or "").lower()
    return (
        alignment.startswith("right")
        and "\n" in block.text
        and re.search(r"(?:19|20)\d{2}\s*年", block.text) is not None
    )
